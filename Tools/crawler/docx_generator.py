import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def save_docx(output_path, title, skill, passage_text, questions):
    doc = Document()
    
    # Configure styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{title} ({skill.upper()})")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Passage / Script Section
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("PART 1: PASSAGE / TRANSCRIPT")
    h1_run.font.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.color.rgb = RGBColor(15, 23, 42)
    
    p_text = doc.add_paragraph(passage_text)
    p_text.paragraph_format.space_after = Pt(24)
    
    # Questions Section
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("PART 2: QUESTIONS")
    h2_run.font.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.color.rgb = RGBColor(15, 23, 42)
    
    for q in questions:
        q_id = q['id']
        q_text = q['text']
        q_type = q['type']
        
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(12)
        q_run = qp.add_run(f"Question {q_id}: {q_text}")
        q_run.font.bold = True
        
        if q_type == 'fill-in-the-blank':
            doc.add_paragraph("Answer: _______________________")
        elif q_type in ['single-choice', 'multiple-choice']:
            for idx, opt in enumerate(q.get('options', [])):
                doc.add_paragraph(f"[ ] {opt}", style='List Bullet')
                
    # Page Break for Answer Key
    doc.add_page_break()
    
    # Answers & Explanations Section
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("PART 3: ANSWER KEY & EXPLANATIONS")
    h3_run.font.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.color.rgb = RGBColor(16, 185, 129)
    
    for q in questions:
        q_id = q['id']
        q_text = q['text']
        correct = q['correct_answer']
        explanation = q.get('explanation', 'No explanation provided.')
        
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(12)
        a_run = ap.add_run(f"Question {q_id} Answer: ")
        a_run.font.bold = True
        
        display_ans = ", ".join(correct) if isinstance(correct, list) else str(correct)
        ans_val_run = ap.add_run(display_ans)
        ans_val_run.font.bold = True
        ans_val_run.font.color.rgb = RGBColor(16, 185, 129)
        
        exp_p = doc.add_paragraph()
        exp_run = exp_p.add_run(f"Explanation: {explanation}")
        exp_run.font.italic = True
        exp_run.font.color.rgb = RGBColor(100, 116, 139)
        
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Generated DOCX exercise: {output_path}")
