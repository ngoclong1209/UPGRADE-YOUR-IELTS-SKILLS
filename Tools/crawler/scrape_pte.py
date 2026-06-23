import os
import re
import time
import base64
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Target URLs and directory paths
DATA_DIR = "/Users/vungoclong/Desktop/Antigravity/UPGRADE YOUR ILETS SKILLS/practicepteonline"
LISTENING_INDEX_URL = "https://practicepteonline.com/ielts-listening-tests-ielts-listening-practice-test-ielts-listening-pdf-ielts-cambridge-test/"
READING_INDEX_URL = "https://practicepteonline.com/ielts-reading-tests/"

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
}

def download_image(img_url, dest_folder, idx):
    if img_url.startswith('/'):
        img_url = "https://practicepteonline.com" + img_url
    
    # Try to extract extension, default to png
    ext = ".png"
    if '.' in img_url.split('/')[-1]:
        possible_ext = '.' + img_url.split('/')[-1].split('.')[-1]
        possible_ext = possible_ext.split('?')[0]
        if 2 <= len(possible_ext) <= 5:
            ext = possible_ext
            
    img_name = f"image_{idx}{ext}"
    local_path = os.path.join(dest_folder, img_name)
    os.makedirs(dest_folder, exist_ok=True)
    
    if not os.path.exists(local_path):
        try:
            r = requests.get(img_url, headers=HEADERS, timeout=20)
            if r.status_code == 200:
                with open(local_path, 'wb') as f:
                    f.write(r.content)
                print(f"Downloaded image: {img_url} -> {local_path}")
            else:
                print(f"Failed to download image {img_url}. Status: {r.status_code}")
                return None
        except Exception as e:
            print(f"Error downloading image {img_url}: {e}")
            return None
def clean_html_content(content_div):
    # Remove script, style, noscript, ins tags
    for el in content_div.find_all(['script', 'style', 'noscript', 'ins']):
        el.decompose()
        
    # Remove ads/social share buttons by common class and id keywords
    noise_keywords = ['addtoany', 'sharedaddy', 'rating', 'author', 'meta', 'adsbygoogle', 'social', 'button', 'comment', 'subscribe', 'wp-block-buttons']
    for el in content_div.find_all(True):
        if not hasattr(el, 'attrs') or el.attrs is None:
            continue
        classes = el.get('class', [])
        if any(any(k in str(c).lower() for k in noise_keywords) for c in classes):
            el.decompose()
            continue
            
        el_id = el.get('id', '')
        if any(k in el_id.lower() for k in noise_keywords):
            if not el_id.startswith('bg-showmore-hidden'):
                el.decompose()

def add_runs_from_html(element, p_docx, test_folder, bold=False, italic=False, underline=False):
    for child in element.children:
        if isinstance(child, str):
            val = str(child)
            if val:
                run = p_docx.add_run(val)
                if bold: run.font.bold = True
                if italic: run.font.italic = True
                if underline: run.font.underline = True
        else:
            name = child.name
            if name in [None, 'script', 'style', 'noscript']:
                continue
            
            next_bold = bold or name in ['strong', 'b']
            next_italic = italic or name in ['em', 'i']
            next_underline = underline or name in ['u']
            
            if name == 'input':
                t = child.get('type', 'text')
                if t == 'checkbox':
                    run = p_docx.add_run('[ ] ')
                else:
                    run = p_docx.add_run(' _______ ')
                if next_bold: run.font.bold = True
            elif name == 'img':
                src = child.get('src')
                if src:
                    local_path = os.path.join(test_folder, src)
                    if os.path.exists(local_path):
                        in_table = child.find_parent('table') is not None
                        w = Inches(1.8) if in_table else Inches(5.5)
                        try:
                            run = p_docx.add_run()
                            run.add_picture(local_path, width=w)
                        except Exception as e:
                            print(f"Error embedding picture in docx: {e}")
            elif name == 'br':
                p_docx.add_run('\n')
            elif name in ['span', 'strong', 'b', 'em', 'i', 'u', 'a']:
                add_runs_from_html(child, p_docx, test_folder, next_bold, next_italic, next_underline)
            else:
                txt = child.get_text()
                if txt:
                    run = p_docx.add_run(txt)
                    if next_bold: run.font.bold = True
                    if next_italic: run.font.italic = True
                    if next_underline: run.font.underline = True

def html_table_to_docx(table_element, doc, test_folder):
    rows = table_element.find_all('tr')
    if not rows:
        return
    max_cols = 0
    for r in rows:
        cells = r.find_all(['td', 'th'])
        max_cols = max(max_cols, len(cells))
    if max_cols == 0:
        return
    
    t_docx = doc.add_table(rows=len(rows), cols=max_cols)
    t_docx.style = 'Table Grid'
    
    for row_idx, r in enumerate(rows):
        cells = r.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            cell_docx = t_docx.rows[row_idx].cells[col_idx]
            p_cell = cell_docx.paragraphs[0]
            add_runs_from_html(cell, p_cell, test_folder)

def get_test_links(index_url, pattern_word):
    print(f"Fetching index: {index_url}")
    try:
        r = requests.get(index_url, headers=HEADERS, timeout=15)
        if r.status_code != 200:
            print(f"Failed to fetch index page. Status code: {r.status_code}")
            return []
        
        soup = BeautifulSoup(r.text, 'html.parser')
        links = []
        for a in soup.find_all('a', href=True):
            href = a['href']
            text = a.get_text().strip()
            if href.startswith('/'):
                href = "https://practicepteonline.com" + href
            
            if pattern_word in href.lower() and re.search(r'-test-\d+/?$', href.lower()):
                links.append((href, text))
            elif pattern_word in href.lower() and re.search(r'-test-\d+-', href.lower()):
                links.append((href, text))
            elif 'listening' in pattern_word and ('listening' in href.lower() or 'test' in href.lower()) and re.search(r'-(\d+)/?$', href):
                links.append((href, text))
        
        seen = set()
        unique_links = []
        for href, text in links:
            if href not in seen:
                seen.add(href)
                unique_links.append((href, text))
        
        print(f"Found {len(unique_links)} unique {pattern_word} tests.")
        return unique_links
    except Exception as e:
        print(f"Error fetching index {index_url}: {e}")
        return []

def extract_answers_text(html, soup):
    btn = soup.find('button', id=re.compile(r'^bg-showmore-action-'))
    if not btn:
        return ""
    
    btn_id = btn.get('id')
    ans_id = btn_id.replace('bg-showmore-action-', 'bg-showmore-hidden-')
    
    m = re.search(r"id=['\"]" + re.escape(ans_id) + r"['\"][^>]*>(.*?)(?:</div>|<ins|class=\"addtoany)", html, re.DOTALL | re.IGNORECASE)
    block = ""
    if m:
        block = m.group(1)
    else:
        ans_div = soup.find(id=ans_id)
        if ans_div:
            block = str(ans_div)
            
    if not block:
        return ""
        
    clean_soup = BeautifulSoup(block, 'html.parser')
    lines = []
    
    lists = clean_soup.find_all(['ol', 'ul'])
    if lists:
        for lst in lists:
            if lst.name == 'ol':
                idx = 1
                for li in lst.find_all('li'):
                    li_text = li.get_text().strip()
                    if li_text:
                        if not re.match(r'^\d+[\.\s\-]', li_text):
                            lines.append(f"{idx}. {li_text}")
                        else:
                            lines.append(li_text)
                        idx += 1
            else:
                for li in lst.find_all('li'):
                    li_text = li.get_text().strip()
                    if li_text:
                        lines.append(f"- {li_text}")
    else:
        block_text = block
        block_text = re.sub(r'</li>', '\n', block_text)
        block_text = re.sub(r'</p>', '\n', block_text)
        block_text = re.sub(r'<br\s*/?>', '\n', block_text)
        block_text = re.sub(r'</td>', '\n', block_text)
        block_text = re.sub(r'</tr>', '\n', block_text)
        block_text = re.sub(r'</div>', '\n', block_text)
        
        temp_soup = BeautifulSoup(block_text, 'html.parser')
        text = temp_soup.get_text().strip()
        for line in text.split('\n'):
            line_str = line.strip()
            if line_str:
                lines.append(line_str)
                
    return "\n".join(lines)

def parse_answers_list(answers_text):
    # Parses ANSWER KEY text lines into a list of pure string answers
    answers = []
    for line in answers_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Check if line matches question format e.g. "1. yes" or "2. no"
        m = re.match(r'^\d+[\.\s\-]+(.*)', line)
        if m:
            answers.append(m.group(1).strip())
        else:
            # If not matching list, just add the line
            answers.append(line)
    return answers

def get_image_base64(file_path):
    ext = file_path.split('.')[-1].lower()
    mime_type = f"image/{ext}"
    if ext in ['jpg', 'jpeg']:
        mime_type = "image/jpeg"
    elif ext == 'svg':
        mime_type = "image/svg+xml"
    try:
        with open(file_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode('utf-8')
            return f"data:{mime_type};base64,{encoded}"
    except Exception as e:
        print(f"Error encoding image {file_path}: {e}")
        return ""

def save_interactive_html(output_path, title, test_type, idx_num, content_html, answers, audio_filename=None):
    test_folder = os.path.dirname(output_path)
    # Convert local images inside content_html into base64 embedded data
    content_soup = BeautifulSoup(content_html, 'html.parser')
    for img in content_soup.find_all('img'):
        src = img.get('src')
        if src and not src.startswith('data:'):
            local_path = os.path.join(test_folder, src)
            if os.path.exists(local_path):
                base64_data = get_image_base64(local_path)
                if base64_data:
                    img['src'] = base64_data
    content_html_clean = str(content_soup)

    audio_element = ""
    if audio_filename:
        audio_element = f'''
        <div class="audio-fixed-bottom">
            <audio src="{audio_filename}" controls></audio>
        </div>
        '''
        
    answer_rows = []
    for idx, correct in enumerate(answers):
        q_num = idx + 1
        row = f'''
        <div class="answer-row">
            <span class="q-number">{q_num}</span>
            <input type="text" class="q-input" id="q-{q_num}" placeholder="Type answer...">
            <span class="feedback-msg" id="feedback-{q_num}"></span>
        </div>
        '''
        answer_rows.append(row)
        
    rows_html = "\n".join(answer_rows)
    
    html_template = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-primary: #ffffff;
            --bg-secondary: #f9fafb;
            --text-primary: #111827;
            --text-secondary: #4b5563;
            --border-color: #e5e7eb;
            --accent: #111827;
            --accent-hover: #1f2937;
            --success: #059669;
            --error: #dc2626;
            --success-bg: #ecfdf5;
            --error-bg: #fef2f2;
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
            font-family: 'Plus Jakarta Sans', sans-serif;
        }}

        body {{
            background-color: var(--bg-primary);
            color: var(--text-primary);
            display: flex;
            flex-direction: column;
            height: 100vh;
            overflow: hidden;
        }}

        header {{
            background-color: var(--bg-primary);
            border-bottom: 1px solid var(--border-color);
            padding: 1rem 2rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            height: 70px;
            flex-shrink: 0;
        }}

        h1 {{
            font-size: 1.4rem;
            font-weight: 700;
        }}

        .main-layout {{
            display: flex;
            flex: 1;
            overflow: hidden;
            position: relative;
        }}

        .passage-pane {{
            flex: 2;
            overflow-y: auto;
            padding: 2.5rem 3.5rem;
            border-right: 1px solid var(--border-color);
            padding-bottom: 110px;
            position: relative;
        }}

        .answer-pane {{
            flex: 1;
            overflow-y: auto;
            padding: 2.5rem 2rem;
            background-color: var(--bg-secondary);
            display: flex;
            flex-direction: column;
            flex-shrink: 0;
            width: 380px;
            padding-bottom: 110px;
        }}

        .audio-fixed-bottom {{
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            height: 80px;
            background: #ffffff;
            border-top: 1px solid var(--border-color);
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 10px 2rem;
            z-index: 1000;
            box-shadow: 0 -4px 12px rgba(0, 0, 0, 0.05);
        }}

        .audio-fixed-bottom audio {{
            width: 100%;
            max-width: 800px;
        }}

        .content-html {{
            font-size: 1.05rem;
            line-height: 1.8;
        }}

        .content-html p {{
            margin-bottom: 1rem;
        }}

        .content-html table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5rem 0;
        }}

        .content-html th, .content-html td {{
            border: 1px solid var(--border-color);
            padding: 0.75rem;
            text-align: left;
        }}

        .content-html img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5rem auto;
            border-radius: 8px;
        }}

        .content-html button, .content-html [id^="bg-showmore-hidden"] {{
            display: none !important;
        }}

        .sheet-title {{
            font-size: 1.3rem;
            font-weight: 700;
            margin-bottom: 1.5rem;
            border-bottom: 2px solid var(--border-color);
            padding-bottom: 0.5rem;
        }}

        .answer-grid {{
            display: flex;
            flex-direction: column;
            gap: 1rem;
            margin-bottom: 2rem;
        }}

        .answer-row {{
            display: flex;
            align-items: center;
            gap: 1rem;
            background: var(--bg-primary);
            padding: 0.75rem 1rem;
            border-radius: 8px;
            border: 1px solid var(--border-color);
        }}

        .q-number {{
            font-weight: 700;
            color: var(--accent);
            width: 30px;
        }}

        .q-input {{
            flex: 1;
            background: var(--bg-secondary);
            border: 1px solid var(--border-color);
            color: var(--text-primary);
            padding: 0.5rem 0.75rem;
            border-radius: 6px;
            font-size: 1rem;
            outline: none;
        }}

        .q-input:focus {{
            border-color: var(--accent);
        }}

        .feedback-msg {{
            font-size: 0.9rem;
            display: none;
            font-weight: 600;
        }}

        .feedback-correct {{
            border-color: var(--success) !important;
            background-color: var(--success-bg) !important;
            color: var(--success) !important;
        }}

        .feedback-incorrect {{
            border-color: var(--error) !important;
            background-color: var(--error-bg) !important;
            color: var(--error) !important;
        }}

        .check-btn {{
            background-color: var(--accent);
            color: white;
            border: none;
            padding: 1rem;
            font-size: 1.1rem;
            font-weight: 700;
            border-radius: 8px;
            cursor: pointer;
            transition: background-color 0.2s;
            width: 100%;
        }}

        .check-btn:hover {{
            background-color: var(--accent-hover);
        }}

        .score-display {{
            display: none;
            margin-top: 1.5rem;
            padding: 1rem;
            background: var(--success-bg);
            border: 1px solid var(--success);
            border-radius: 8px;
            text-align: center;
            font-size: 1.2rem;
            font-weight: 700;
            color: var(--success);
        }}

        /* Highlight tool styling */
        .text-highlight {{
            background-color: #fef08a !important;
            color: #000000 !important;
            cursor: pointer;
        }}
        
        .highlight-floating-btn {{
            position: absolute;
            display: none;
            background-color: #111827;
            color: #ffffff;
            border: none;
            padding: 6px 12px;
            border-radius: 6px;
            font-size: 0.85rem;
            font-weight: 600;
            cursor: pointer;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
            z-index: 10000;
        }}
        
        .highlight-floating-btn:hover {{
            background-color: #1f2937;
        }}

        @media (max-width: 900px) {{
            body {{
                overflow: auto;
                height: auto;
            }}
            .main-layout {{
                flex-direction: column;
                height: auto;
                overflow: visible;
            }}
            .passage-pane {{
                border-right: none;
                border-bottom: 1px solid var(--border-color);
                padding-bottom: 2rem;
            }}
            .answer-pane {{
                width: 100%;
                padding-bottom: 100px;
            }}
        }}
    </style>
</head>
<body>
    <header>
        <h1>{title}</h1>
    </header>

    <div class="main-layout">
        <button class="highlight-floating-btn" id="float-hl-btn">Highlight</button>
        <div class="passage-pane">
            <div class="content-html">
                {content_html_clean}
            </div>
        </div>
        
        <div class="answer-pane">
            <div class="sheet-title">Answer Sheet</div>
            <div class="answer-grid">
                {rows_html}
            </div>
            <button class="check-btn" onclick="checkAnswers()">Check Answers</button>
            <div class="score-display" id="score-box"></div>
        </div>
        
        {audio_element}
    </div>

    <script>
        const answers = {json.dumps(answers, ensure_ascii=False)};

        function checkAnswers() {{
            let score = 0;
            answers.forEach((correct, index) => {{
                const qNum = index + 1;
                const inputEl = document.getElementById(`q-${{qNum}}`);
                const feedbackEl = document.getElementById(`feedback-${{qNum}}`);
                
                if (inputEl) {{
                    const userVal = inputEl.value.trim().toLowerCase();
                    const isCorrect = (userVal === correct.trim().toLowerCase());
                    
                    if (isCorrect) {{
                        score++;
                        inputEl.className = 'q-input feedback-correct';
                        feedbackEl.style.display = 'inline';
                        feedbackEl.innerHTML = '✓ Correct';
                    }} else {{
                        inputEl.className = 'q-input feedback-incorrect';
                        feedbackEl.style.display = 'inline';
                        feedbackEl.innerHTML = `✗ ${{correct}}`;
                    }}
                }}
            }});

            const scoreBox = document.getElementById('score-box');
            scoreBox.style.display = 'block';
            scoreBox.innerHTML = `Your Score: ${{score}} / ${{answers.length}} (${{Math.round((score / answers.length) * 100)}}%)`;
        }}

        // Text Highlight Engine
        const hlBtn = document.getElementById('float-hl-btn');
        const passagePane = document.querySelector('.passage-pane');

        document.addEventListener('selectionchange', () => {{
            const selection = window.getSelection();
            if (selection.rangeCount > 0 && selection.toString().trim().length > 0) {{
                const range = selection.getRangeAt(0);
                if (passagePane.contains(range.commonAncestorContainer)) {{
                    const rect = range.getBoundingClientRect();
                    hlBtn.style.display = 'block';
                    hlBtn.style.top = `${{window.scrollY + rect.top - 40}}px`;
                    hlBtn.style.left = `${{window.scrollX + rect.left + (rect.width/2) - 45}}px`;
                    return;
                }}
            }}
            hlBtn.style.display = 'none';
        }});

        hlBtn.addEventListener('mousedown', (e) => {{
            e.preventDefault();
            highlightSelection();
            hlBtn.style.display = 'none';
        }});

        function highlightSelection() {{
            const selection = window.getSelection();
            if (selection.rangeCount > 0) {{
                const range = selection.getRangeAt(0);
                const span = document.createElement("span");
                span.className = "text-highlight";
                span.title = "Click to remove highlight";
                span.onclick = function() {{
                    this.replaceWith(...this.childNodes);
                }};
                try {{
                    range.surroundContents(span);
                }} catch(err) {{
                    console.log("Cross element highlighting failed:", err);
                }}
                selection.removeAllRanges();
            }}
        }}
    </script>
</body>
</html>'''
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_template)
    print(f"Successfully saved interactive HTML: {output_path}")

def process_test(url, test_type, idx_num):
    test_folder = os.path.join(DATA_DIR, test_type, f"Test_{idx_num}")
    docx_path = os.path.join(test_folder, f"Test_{idx_num}.docx")
    html_path = os.path.join(test_folder, f"Test_{idx_num}.html")
    
    # Check if files exist and are upgraded (contains highlight functionality)
    is_upgraded = False
    if os.path.exists(docx_path) and os.path.exists(html_path):
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if 'Text Highlight Engine' in content:
                    is_upgraded = True
        except:
            pass
            
    if is_upgraded:
        print(f"[{test_type.upper()}] Test {idx_num} already exists and is fully upgraded. Skipping.")
        return
    
    print(f"[{test_type.upper()}] Processing Test {idx_num}: {url}")
    try:
        r = requests.get(url, headers=HEADERS, timeout=20)
        if r.status_code != 200:
            print(f"Failed to fetch {url}. Status: {r.status_code}")
            return
        
        soup = BeautifulSoup(r.text, 'html.parser')
        content_div = soup.find(class_='entry-content')
        if not content_div:
            print(f"No entry-content div found for {url}")
            return
            
        # 1. Download and map unique non-noscript images
        img_tags = content_div.find_all('img')
        local_images = {}
        img_idx = 1
        for img in img_tags:
            if img.find_parent('noscript'):
                continue
            src = img.get('data-src') or img.get('src')
            if not src or src.startswith('data:image'):
                continue
            if src not in local_images:
                local_path = download_image(src, os.path.join(test_folder, "images"), img_idx)
                if local_path:
                    local_images[src] = local_path
                    img_idx += 1

        # 2. Update img tags in BS4 content to point to relative paths
        for img in content_div.find_all('img'):
            if img.find_parent('noscript'):
                img.decompose()  # Remove noscript fallbacks
                continue
            src = img.get('data-src') or img.get('src')
            if src in local_images:
                rel_path = os.path.relpath(local_images[src], test_folder)
                img['src'] = rel_path
                if 'data-src' in img.attrs:
                    del img['data-src']
                if 'data-lazyloaded' in img.attrs:
                    del img['data-lazyloaded']

        # 3. Handle Audio Download if Listening
        audio_src = None
        audio_filename = None
        if test_type == "listening":
            audio_tag = soup.find(['audio', 'source'])
            if audio_tag:
                audio_src = audio_tag.get('src')
            if not audio_src:
                mp3_match = re.search(r'href="([^"]+\.mp3)"', r.text)
                if mp3_match:
                    audio_src = mp3_match.group(1)
            
            if audio_src:
                if audio_src.startswith('/'):
                    audio_src = "https://practicepteonline.com" + audio_src
                
                audio_filename = f"audio_{idx_num}.mp3"
                audio_path = os.path.join(test_folder, audio_filename)
                os.makedirs(test_folder, exist_ok=True)
                
                if not os.path.exists(audio_path):
                    print(f"Downloading audio: {audio_src}")
                    ar = requests.get(audio_src, headers=HEADERS, timeout=30)
                    if ar.status_code == 200:
                        with open(audio_path, 'wb') as f:
                            f.write(ar.content)
                        print(f"Audio downloaded to {audio_path}")
                    else:
                        print(f"Failed to download audio. Status: {ar.status_code}")
        
        # 4. Extract Answer Key
        answers_txt = extract_answers_text(r.text, soup)
        answers_list = parse_answers_list(answers_txt)
        
        # Clean up answer widgets and metadata/ads from content_div
        for btn in content_div.find_all('button', id=re.compile(r'^bg-showmore-action-')):
            btn.decompose()
        for hidden in content_div.find_all(id=re.compile(r'^bg-showmore-hidden-')):
            hidden.decompose()
        clean_html_content(content_div)
        
        # 5. Create DOCX
        doc = Document()
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        
        # Document Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"IELTS {test_type.capitalize()} Practice - Test {idx_num}")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(79, 70, 229)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # Parse and write content to DOCX
        for child in content_div.find_all(recursive=False):
            if child.find('button', id=re.compile(r'^bg-showmore-action-')) or (child.name == 'button' and 'bg-showmore-action' in child.get('id', '')):
                break
            if child.get('id', '').startswith('bg-showmore-hidden-') or child.find(id=re.compile(r'^bg-showmore-hidden-')):
                break
            if child.name in ['ins', 'script', 'style'] or 'adsbygoogle' in str(child) or 'addtoany' in str(child.get('class', [])):
                continue
            
            if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(child.name[1])
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(12)
                hp_run = hp.add_run(child.get_text().strip())
                hp_run.font.bold = True
                hp_run.font.size = Pt(16 - level)
                hp_run.font.color.rgb = RGBColor(15, 23, 42)
            elif child.name == 'p':
                txt = child.get_text().strip()
                if not txt and not child.find('input') and not child.find('img'):
                    continue
                pp = doc.add_paragraph()
                add_runs_from_html(child, pp, test_folder)
            elif child.name == 'figure':
                table_el = child.find('table')
                if table_el:
                    html_table_to_docx(table_el, doc, test_folder)
                else:
                    caption = child.find('figcaption')
                    if caption:
                        cp = doc.add_paragraph()
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        c_run = cp.add_run(caption.get_text().strip())
                        c_run.font.italic = True
                        c_run.font.size = Pt(9.5)
            elif child.name == 'table':
                html_table_to_docx(child, doc, test_folder)
            elif child.name in ['ul', 'ol']:
                for li in child.find_all('li'):
                    lp = doc.add_paragraph(style='List Bullet')
                    add_runs_from_html(li, lp, test_folder)
        
        # Page Break & Answer Key for DOCX
        doc.add_page_break()
        ans_title_p = doc.add_paragraph()
        ans_title_run = ans_title_p.add_run("ANSWER KEY")
        ans_title_run.font.bold = True
        ans_title_run.font.size = Pt(14)
        ans_title_run.font.color.rgb = RGBColor(16, 185, 129)
        ans_title_p.paragraph_format.space_after = Pt(12)
        
        if answers_txt:
            for line in answers_txt.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("Answer key was not found or could not be extracted.")
        
        doc.save(docx_path)
        print(f"Successfully saved Word file: {docx_path}")
        
        # 6. Create Interactive HTML
        title_str = f"IELTS {test_type.capitalize()} Practice - Test {idx_num}"
        save_interactive_html(html_path, title_str, test_type, idx_num, str(content_div), answers_list, audio_filename)
        
    except Exception as e:
        print(f"Error processing {url}: {e}")
        import traceback
        traceback.print_exc()

def main():
    listening_links = get_test_links(LISTENING_INDEX_URL, "listening")
    listening_tests = []
    for href, text in listening_links:
        num_match = re.search(r'test-(\d+)/?$', href.lower())
        if not num_match:
            num_match = re.search(r'listening-(\d+)/?$', href.lower())
        if num_match:
            num = int(num_match.group(1))
            listening_tests.append((num, href))
    listening_tests.sort(key=lambda x: x[0])
    
    reading_links = get_test_links(READING_INDEX_URL, "reading")
    reading_tests = []
    for href, text in reading_links:
        num_match = re.search(r'test-(\d+)/?$', href.lower())
        if num_match:
            num = int(num_match.group(1))
            reading_tests.append((num, href))
    reading_tests.sort(key=lambda x: x[0])
    
    print("\n--- Starting Scraping ---")
    print(f"Total Listening Tests: {len(listening_tests)}")
    print(f"Total Reading Tests: {len(reading_tests)}")
    
    # Run all tests sequentially
    # Run all tests sequentially
    for num, href in listening_tests:
        html_path = os.path.join(DATA_DIR, "listening", f"Test_{num}", f"Test_{num}.html")
        if not os.path.exists(html_path):
            print(f"Listening Test {num} missing. Processing...")
            process_test(href, "listening", num)
            time.sleep(2)
        else:
            print(f"Listening Test {num} exists. Skipping.")
        
    for num, href in reading_tests:
        html_path = os.path.join(DATA_DIR, "reading", f"Test_{num}", f"Test_{num}.html")
        if not os.path.exists(html_path):
            print(f"Reading Test {num} missing. Processing...")
            process_test(href, "reading", num)
            time.sleep(2)
        else:
            print(f"Reading Test {num} exists. Skipping.")

if __name__ == "__main__":
    main()
