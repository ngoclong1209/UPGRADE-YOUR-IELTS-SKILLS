import os
import re
import json
from docx import Document

BASE_DIR = "/Users/vungoclong/Desktop/Antigravity/UPGRADE YOUR ILETS SKILLS"
DOCX_DIR = os.path.join(BASE_DIR, "practicepteonline/Listening docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "practicepteonline/listening")
TEMPLATE_PATH = os.path.join(BASE_DIR, "pte-listening-audios/template.html")

def parse_docx_to_html(docx_path):
    doc = Document(docx_path)
    
    # We will split paragraphs into passage content and answers at the end.
    # Typically, answers are listed at the very end as "1. answer", "2. answer".
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append(p.text)
        
    tables = doc.tables
    
    # Let's find the answers index
    # We iterate backwards to find the first paragraph that doesn't look like an answer key item
    answers = {}
    answer_idx_start = len(paragraphs)
    for i in range(len(paragraphs) - 1, -1, -1):
        text = paragraphs[i].strip()
        if not text:
            continue
        # Answer format: "1. 300" or "31. B"
        m = re.match(r'^(\d+)\.\s+(.*)$', text)
        if m:
            num = int(m.group(1))
            answers[num] = m.group(2).strip()
            answer_idx_start = i
        else:
            # If we find a line that is NOT an answer (and not empty), we stop.
            break

    # Reconstruct the passage HTML
    passage_html = []
    
    # To handle tables seamlessly, we'll iterate through the document elements in order.
    # Python-docx doesn't easily interleave paragraphs and tables sequentially using doc.paragraphs and doc.tables.
    # We will use doc.element.body.iterchildren()
    
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    
    current_idx = 0
    for child in doc.element.body.iterchildren():
        if current_idx >= answer_idx_start:
            # Stop if we reached the answers
            break
            
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if text:
                # Replace [ ] with a stylized checkbox if desired, or keep as is.
                # text = text.replace('[ ]', '☐')
                if text.startswith('Part '):
                    passage_html.append(f'<h2>{text}</h2>')
                elif text.startswith('Questions '):
                    passage_html.append(f'<h3>{text}</h3>')
                else:
                    passage_html.append(f'<p>{text}</p>')
            current_idx += 1
            
        elif isinstance(child, CT_Tbl):
            t = Table(child, doc)
            passage_html.append('<table border="1" style="width:100%; border-collapse: collapse; margin-bottom: 1rem;">')
            for row in t.rows:
                passage_html.append('<tr>')
                for cell in row.cells:
                    passage_html.append(f'<td style="padding: 5px;">{cell.text.strip()}</td>')
                passage_html.append('</tr>')
            passage_html.append('</table>')

    # Ensure answers list is sorted and continuous
    if not answers:
        print(f"Warning: No answers found in {docx_path}")
        return "", [], ""
        
    max_ans = max(answers.keys())
    answers_list = [answers.get(i, "") for i in range(1, max_ans + 1)]
    
    return "\n".join(passage_html), answers_list, paragraphs[0] if paragraphs else "Listening Test"

def build_html_for_test(idx_num):
    docx_path = os.path.join(DOCX_DIR, f"Test_{idx_num}.docx")
    if not os.path.exists(docx_path):
        return False
        
    print(f"Building Test_{idx_num}...")
    passage_html, answers, title = parse_docx_to_html(docx_path)
    
    with open(TEMPLATE_PATH, 'r', encoding='utf-8') as f:
        template = f.read()
        
    # Build Answer Pane
    num_questions = len(answers)
    parts = []
    if num_questions <= 4:
        parts = [(1, num_questions)]
    else:
        base_size = num_questions // 4
        rem = num_questions % 4
        start = 1
        for i in range(4):
            size = base_size + (1 if i < rem else 0)
            end = start + size - 1
            if size > 0:
                parts.append((start, end))
                start = end + 1

    answer_rows = []
    for q_idx, correct in enumerate(answers):
        q_num = q_idx + 1
        part_idx = -1
        part_end = -1
        for p_idx, (p_start, p_end) in enumerate(parts):
            if q_num == p_start:
                part_idx = p_idx + 1
                part_end = p_end
                break
                
        if part_idx != -1:
            if part_idx > 1:
                answer_rows.append('</div>')
            display_style = "block" if part_idx == 1 else "none"
            answer_rows.append(f'<div class="question-part-section" id="part-section-{part_idx}" style="display: {display_style}; border-left: 3px solid var(--accent); padding-left: 10px; margin-bottom: 15px;">')
            answer_rows.append(f'''
            <div class="part-header-container">
                <span class="part-header-title">✨ PART {part_idx} (Q{q_num}-{part_end}) ✨</span>
                <button class="part-check-btn" onclick="checkPart({q_num}, {part_end})">Check Part {part_idx}</button>
            </div>
            ''')
            
        row = f'''
        <div class="answer-row" style="position: relative; display: flex; align-items: center;">
            <span class="q-number">🧸 Q{q_num}</span>
            <input type="text" class="q-input" id="q-{q_num}" placeholder="Your answer...">
            <button class="strategy-hint-btn" onclick="showStrategy({q_num}, event)" title="💡 IELTS 9.0 Strategy">💡</button>
            <span class="feedback-msg" id="feedback-{q_num}"></span>
        </div>
        '''
        answer_rows.append(row)
        
    if len(parts) > 0:
        answer_rows.append('</div>')
        
    answer_html = "\n".join(answer_rows)
    
    # Inject into template
    # Replace the passage placeholder
    soup_html = template.replace('{{TITLE}}', f"IELTS Listening Practice - Test {idx_num}")
    
    # We need to inject passage_html inside <div class="passage-pane"> and answer_html inside <div class="answer-grid">
    # We will use simple string replacement or regex
    
    import bs4
    soup = bs4.BeautifulSoup(soup_html, 'html.parser')
    
    passage_pane = soup.find(class_='passage-pane')
    if passage_pane:
        passage_pane.clear()
        passage_pane.append(bs4.BeautifulSoup(passage_html, 'html.parser'))
        
    answer_grid = soup.find(class_='answer-grid')
    if answer_grid:
        answer_grid.clear()
        answer_grid.append(bs4.BeautifulSoup(answer_html, 'html.parser'))
        
    # Inject answers array into script
    # We can append a script at the end of body
    script_content = f"const answers = {json.dumps(answers)};\n"
    new_script = soup.new_tag("script")
    new_script.string = script_content
    soup.body.append(new_script)
    
    # Set Audio SRC
    audio_player = soup.find('audio')
    if audio_player:
        audio_player['src'] = f"https://cdn.jsdelivr.net/gh/ngoclong1209/practicepteonline/Listening_audios/Test_{idx_num}.mp3"
        audio_container = soup.find(id='audio-player-container')
        if audio_container:
            audio_container['style'] = "display: flex;"
    
    test_out_dir = os.path.join(OUTPUT_DIR, f"Test_{idx_num}")
    os.makedirs(test_out_dir, exist_ok=True)
    
    out_html = os.path.join(test_out_dir, "Test_{idx_num}.html")
    # Write to Test_X.html (as we renamed index.html to Test_X.html previously)
    with open(out_html, 'w', encoding='utf-8') as f:
        f.write(str(soup))
        
    # Also write index.html for compatibility if needed? No, user repo uses Test_X.html now. Let's just create Test_X.html
    # Wait, github pages needs index.html to work without specifying file name in URL!
    # "bạn kiểm tra lại gói full test trên github vì nó báo lỗi nhé: 404 File not found For root URLs you must provide an index.html file."
    # So we should write BOTH `Test_X.html` and `index.html` (identical or symlink) to ensure it works!
    with open(os.path.join(test_out_dir, "index.html"), 'w', encoding='utf-8') as f:
        f.write(str(soup))
        
    print(f"Test {idx_num} generated.")
    return True

if __name__ == '__main__':
    for i in range(1, 101):
        build_html_for_test(i)
