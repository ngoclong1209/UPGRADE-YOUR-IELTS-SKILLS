import os
import re
import json
import base64
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_DIR = "/Users/vungoclong/Desktop/Antigravity/UPGRADE YOUR ILETS SKILLS/practicepteonline"
# Configurable CDN Base URL for hosting audio files online.
# If left empty, it will fall back to using relative paths (e.g. "audio_X.mp3") for local offline use.
AUDIO_CDN_BASE_URL = "https://cdn.jsdelivr.net/gh/ngoclong1209/pte-listening-audios/" 
# Deployed Google Apps Script URL for student authorization and results tracking
GOOGLE_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbwpA-N8yVVDrmFD4pZRkDpwOGqDJYLA_LSwh0WVrnL6rPxw4ionhhEAfV2b0df_hJaH/exec" 

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
}

def get_sfx_base64_data():
    sfx_files = {
        "click": "clicksfx.mp3",
        "correct": "correctanswersfx.mp3",
        "wrong": "wrongbiteanswersfx.mp3",
        "win": "winningsfx.mp3"
    }
    encoded_sfx = {}
    for key, name in sfx_files.items():
        path = f"/Users/vungoclong/Desktop/XCode/Data R2R/AudioResources/audiosfx/{name}"
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode('utf-8')
                    encoded_sfx[key] = f"data:audio/mp3;base64,{encoded}"
            except Exception as e:
                print(f"Error encoding SFX {name}: {e}")
                encoded_sfx[key] = ""
        else:
            print(f"Warning: SFX path not found: {path}")
            encoded_sfx[key] = ""
    return encoded_sfx

SFX_DATA = get_sfx_base64_data()

def ensure_png_if_webp(local_path):
    if not local_path or not os.path.exists(local_path):
        return local_path
    if local_path.lower().endswith('.webp'):
        png_path = local_path[:-5] + ".png"
        if not os.path.exists(png_path):
            try:
                from PIL import Image
                with Image.open(local_path) as img:
                    img.save(png_path, "PNG")
                print(f"Converted webp to png: {local_path} -> {png_path}")
            except Exception as e:
                print(f"Error converting webp to png: {e}")
                return local_path
        return png_path
    return local_path

def get_image_base64(file_path):
    file_path = ensure_png_if_webp(file_path)
    ext = file_path.split('.')[-1].lower()
    mime_type = f"image/{ext}"
    if ext in ['jpg', 'jpeg']:
        mime_type = "image/jpeg"
    elif ext == 'svg':
        mime_type = "image/svg+xml"
    try:
        with open(file_path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode('utf-8')
            return f"data:{mime_type};base64,{encoded}"
    except Exception as e:
        print(f"Error encoding image {file_path}: {e}")
        return ""

def download_image(img_url, dest_folder, idx):
    if img_url.startswith('/'):
        img_url = "https://practicepteonline.com" + img_url
    
    ext = ".png"
    if '.' in img_url.split('/')[-1]:
        possible_ext = '.' + img_url.split('/')[-1].split('.')[-1]
        possible_ext = possible_ext.split('?')[0]
        if 2 <= len(possible_ext) <= 5:
            ext = possible_ext
            
    img_name = f"image_{idx}{ext}"
    local_path = os.path.join(dest_folder, img_name)
    os.makedirs(dest_folder, exist_ok=True)
    
    if not os.path.exists(local_path):
        try:
            r = requests.get(img_url, headers=HEADERS, timeout=20)
            if r.status_code == 200:
                with open(local_path, 'wb') as f:
                    f.write(r.content)
                print(f"Downloaded image: {img_url} -> {local_path}")
                local_path = ensure_png_if_webp(local_path)
            else:
                print(f"Failed to download image {img_url}. Status: {r.status_code}")
                return None
        except Exception as e:
            print(f"Error downloading image {img_url}: {e}")
            return None
    return local_path

def add_runs_from_html(element, p_docx, test_folder, bold=False, italic=False, underline=False):
    for child in element.children:
        if isinstance(child, str):
            val = str(child)
            if val:
                run = p_docx.add_run(val)
                if bold: run.font.bold = True
                if italic: run.font.italic = True
                if underline: run.font.underline = True
        else:
            name = child.name
            if name in [None, 'script', 'style', 'noscript']:
                continue
            
            next_bold = bold or name in ['strong', 'b']
            next_italic = italic or name in ['em', 'i']
            next_underline = underline or name in ['u']
            
            if name == 'input':
                t = child.get('type', 'text')
                if t == 'checkbox':
                    run = p_docx.add_run('[ ] ')
                else:
                    run = p_docx.add_run(' _______ ')
                if next_bold: run.font.bold = True
            elif name == 'img':
                img_idx = child.get('__img_index')
                if img_idx:
                    img_dir = os.path.join(test_folder, "images")
                    local_path = None
                    if os.path.exists(img_dir):
                        possible_names = [f"image_{img_idx}.png", f"image_{img_idx}.jpg", f"image_{img_idx}.jpeg", f"image_{img_idx}.gif", f"image_{img_idx}.webp"]
                        for fname in possible_names:
                            p = os.path.join(img_dir, fname)
                            if os.path.exists(p):
                                local_path = p
                                break
                    
                    if local_path:
                        local_path = ensure_png_if_webp(local_path)
                    
                    if local_path and os.path.exists(local_path):
                        in_table = child.find_parent('table') is not None
                        w = Inches(1.8) if in_table else Inches(5.5)
                        try:
                            run = p_docx.add_run()
                            run.add_picture(local_path, width=w)
                        except Exception as e:
                            print(f"Error embedding picture in docx: {e}")
            elif name == 'br':
                p_docx.add_run('\n')
            elif name in ['span', 'strong', 'b', 'em', 'i', 'u', 'a']:
                add_runs_from_html(child, p_docx, test_folder, next_bold, next_italic, next_underline)
            else:
                txt = child.get_text()
                if txt:
                    run = p_docx.add_run(txt)
                    if next_bold: run.font.bold = True
                    if next_italic: run.font.italic = True
                    if next_underline: run.font.underline = True

def html_table_to_docx(table_element, doc, test_folder):
    rows = table_element.find_all('tr')
    if not rows:
        return
    max_cols = 0
    for r in rows:
        cells = r.find_all(['td', 'th'])
        max_cols = max(max_cols, len(cells))
    if max_cols == 0:
        return
    
    t_docx = doc.add_table(rows=len(rows), cols=max_cols)
    t_docx.style = 'Table Grid'
    
    for row_idx, r in enumerate(rows):
        cells = r.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            cell_docx = t_docx.rows[row_idx].cells[col_idx]
            p_cell = cell_docx.paragraphs[0]
            add_runs_from_html(cell, p_cell, test_folder)

def parse_element_to_docx(element, doc, test_folder):
    for child in element.children:
        if isinstance(child, str):
            if child.strip():
                p = doc.add_paragraph()
                p.add_run(child.strip())
        else:
            name = child.name
            if name in ['script', 'style', 'noscript', 'button']:
                continue
            
            if name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(name[1])
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(12)
                hp_run = hp.add_run(child.get_text().strip())
                hp_run.font.bold = True
                hp_run.font.size = Pt(16 - level)
                hp_run.font.color.rgb = RGBColor(15, 23, 42)
            elif name == 'table':
                html_table_to_docx(child, doc, test_folder)
            elif name in ['ul', 'ol']:
                for li in child.find_all('li'):
                    lp = doc.add_paragraph(style='List Bullet')
                    add_runs_from_html(li, lp, test_folder)
            elif name == 'p':
                txt = child.get_text().strip()
                if txt or child.find('input') or child.find('img'):
                    p = doc.add_paragraph()
                    add_runs_from_html(child, p, test_folder)
            elif name in ['div', 'figure', 'section', 'article']:
                block_children = [c for c in child.children if not isinstance(c, str) and c.name in ['p', 'table', 'ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'figure', 'section', 'article']]
                if block_children:
                    parse_element_to_docx(child, doc, test_folder)
                else:
                    txt = child.get_text().strip()
                    if txt or child.find('input') or child.find('img'):
                        p = doc.add_paragraph()
                        add_runs_from_html(child, p, test_folder)

def rebuild_files_for_test(test_folder, test_type, idx_num):
    html_path = os.path.join(test_folder, "index.html")
    docx_path = os.path.join(test_folder, f"Test_{idx_num}.docx")
    test_name = f"{test_type.capitalize()} Test_{idx_num}"
    
    if not os.path.exists(html_path):
        print(f"HTML file not found: {html_path}. Skipping.")
        return
        
    print(f"Upgrading: {html_path}")
    
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
        
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 1. Extract answers array from script
    answers = []
    script_tags = soup.find_all('script')
    for script in script_tags:
        if 'const answers = ' in script.text:
            m = re.search(r'const answers\s*=\s*(\[.*?\]);', script.text, re.DOTALL)
            if m:
                try:
                    answers = json.loads(m.group(1))
                except Exception as e:
                    print(f"Error parsing answers json: {e}")
            break
            
    # 2. Extract title
    title_tag = soup.find('title')
    title = title_tag.get_text().strip() if title_tag else f"IELTS {test_type.capitalize()} Practice - Test {idx_num}"
    
    # 3. Extract the entry-content div
    entry_content = soup.find(class_='entry-content')
    if not entry_content:
        entry_content = soup.find(class_='content-html')
    if not entry_content:
        entry_content = soup.find(class_='passage-pane')
        
    if not entry_content:
        print(f"Could not find content div in {html_path}")
        return
        
    # Ensure any placeholder images are replaced with actual base64 or relative images
    img_tags = entry_content.find_all('img')
    img_idx = 1
    for img in img_tags:
        if img.find_parent('noscript'):
            img.decompose()
            continue
        src = img.get('src')
        data_src = img.get('data-src')
        
        # Check if the image needs to be downloaded or is already local
        local_path = None
        img_dir = os.path.join(test_folder, "images")
        
        # Determine actual remote source
        remote_src = None
        if data_src and not data_src.startswith('data:'):
            remote_src = data_src
        elif src and not src.startswith('data:'):
            remote_src = src
            
        if remote_src:
            local_path = download_image(remote_src, img_dir, img_idx)
            img_idx += 1
        else:
            # Maybe the image was already downloaded
            possible_names = [f"image_{img_idx}.png", f"image_{img_idx}.jpg", f"image_{img_idx}.jpeg", f"image_{img_idx}.gif", f"image_{img_idx}.webp"]
            for name in possible_names:
                p = os.path.join(img_dir, name)
                if os.path.exists(p):
                    local_path = p
                    img_idx += 1
                    break
        
        if local_path and os.path.exists(local_path):
            local_path = ensure_png_if_webp(local_path)
            base64_data = get_image_base64(local_path)
            if base64_data:
                img['src'] = base64_data
            else:
                img['src'] = os.path.relpath(local_path, test_folder)
            
            if 'data-src' in img.attrs:
                del img['data-src']
            if 'data-lazyloaded' in img.attrs:
                del img['data-lazyloaded']

    all_imgs = [img for img in entry_content.find_all('img') if not img.find_parent('noscript')]
    for idx, img in enumerate(all_imgs):
        img['__img_index'] = str(idx + 1)
                
    for el in entry_content.find_all(['script', 'style', 'noscript', 'ins', 'audio']):
        el.decompose()
        
    # Strip out any plaintext answer keys (e.g., matching "1. yes 2. no 3. yes") embedded in the passage/description
    for tag in entry_content.find_all(['p', 'div', 'span', 'h4', 'h5', 'h6', 'strong']):
        tag_classes = tag.get('class') if tag.attrs else None
        if tag_classes and any(c in tag_classes for c in ['entry-content', 'content-html', 'passage-pane']):
            continue
        text = tag.get_text().strip()
        clean_text = ' '.join(text.split())
        if re.search(r'^1\.\s+.*?2\.\s+.*?3\.\s+', clean_text):
            tag.decompose()
        
    # Load 9.0 strategy recommendations if available, else generate default fallback values
    strategy_path = os.path.join(test_folder, "strategies.json")
    strategies_data = []
    if os.path.exists(strategy_path):
        try:
            with open(strategy_path, 'r', encoding='utf-8') as sf:
                strategies_data = json.load(sf)
        except Exception as e:
            print(f"Error loading strategies.json in {test_folder}: {e}")
            
    while len(strategies_data) < len(answers):
        q_num = len(strategies_data) + 1
        strategies_data.append({
            "q_num": q_num,
            "type_tip": "IELTS 9.0 scorers focus on synonyms and modifying qualifiers (e.g. only, never, some).",
            "scan_target": f"Scan the passage for key terms related to Q{q_num}.",
            "analysis_logic": "Analyze context carefully and eliminate options that contradict the passage meaning."
        })

    # Dynamic part calculation (split into up to 4 parts)
    num_questions = len(answers)
    parts = []
    if num_questions <= 4:
        parts = [(1, num_questions)]
    else:
        base_size = num_questions // 4
        rem = num_questions % 4
        start = 1
        for i in range(4):
            size = base_size + (1 if i < rem else 0)
            end = start + size - 1
            if size > 0:
                parts.append((start, end))
                start = end + 1

    answer_rows = []
    for idx, correct in enumerate(answers):
        q_num = idx + 1
        
        # Check if this question starts a part
        part_idx = -1
        part_end = -1
        for p_idx, (p_start, p_end) in enumerate(parts):
            if q_num == p_start:
                part_idx = p_idx + 1
                part_end = p_end
                break
                
        if part_idx != -1:
            if part_idx > 1:
                answer_rows.append('</div>') # Close previous part wrapper
            
            # Start new part wrapper. Only Part 1 is displayed initially.
            display_style = "block" if part_idx == 1 else "none"
            answer_rows.append(f'<div class="question-part-section" id="part-section-{part_idx}" style="display: {display_style}; border-left: 3px solid var(--accent); padding-left: 10px; margin-bottom: 15px;">')
            
            answer_rows.append(f'''
            <div class="part-header-container">
                <span class="part-header-title">✨ PART {part_idx} (Q{q_num}-{part_end}) ✨</span>
                <button class="part-check-btn" onclick="checkPart({q_num}, {part_end})">Check Part {part_idx}</button>
            </div>
            ''')
            
        row = f'''
        <div class="answer-row" style="position: relative; display: flex; align-items: center;">
            <span class="q-number">🧸 Q{q_num}</span>
            <input type="text" class="q-input" id="q-{q_num}" placeholder="Your answer...">
            <button class="strategy-hint-btn" onclick="showStrategy({q_num}, event)" title="💡 IELTS 9.0 Strategy">💡</button>
            <span class="feedback-msg" id="feedback-{q_num}"></span>
        </div>
        '''
        answer_rows.append(row)
        
    if len(parts) > 0:
        answer_rows.append('</div>') # Close the last part wrapper
        
    rows_html = "\n".join(answer_rows)
    
    audio_element = ""
    audio_filename = f"audio_{idx_num}.mp3"
    audio_path = os.path.join(test_folder, audio_filename)
    if test_type == "listening" and (AUDIO_CDN_BASE_URL or os.path.exists(audio_path)):
        if AUDIO_CDN_BASE_URL:
            # Construct the absolute CDN URL (e.g., https://.../Test_X/audio_X.mp3)
            audio_src = f"{AUDIO_CDN_BASE_URL.rstrip('/')}/Test_{idx_num}/{audio_filename}"
        else:
            # Fallback to local relative path
            audio_src = audio_filename
            
        audio_element = f'''
        <div class="audio-fixed-bottom" id="audio-player-container" style="display: none;">
            <audio src="{audio_src}" controls></audio>
        </div>
        '''
        
    html_template = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <link href="https://fonts.googleapis.com/css2?family=Fredoka:wght@300..700&family=DynaPuff:wght@400..700&family=Quicksand:wght@300..700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-primary: #fff5f7;
            --bg-secondary: #fffafb;
            --text-primary: #5c404d;
            --text-secondary: #8c6d7d;
            --border-color: #ffccd5;
            --accent: #ff85a2;
            --accent-hover: #ff5e7e;
            --accent-lilac: #b19ffb;
            --success: #4ad99d;
            --error: #ff768a;
            --success-bg: #e2f9ee;
            --error-bg: #fff0f2;
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
            font-family: 'Quicksand', sans-serif;
        }}

        body {{
            background: linear-gradient(135deg, #fff0f5 0%, #e6e6fa 50%, #f0f8ff 100%);
            background-size: cover;
            color: var(--text-primary);
            display: flex;
            flex-direction: column;
            height: 100vh;
            overflow: hidden;
            position: relative;
            
            /* Block text selection for anti-copy */
            -webkit-user-select: none;
            -moz-user-select: none;
            -ms-user-select: none;
            user-select: none;
        }}

        input, textarea, [contenteditable="true"] {{
            -webkit-user-select: text;
            -moz-user-select: text;
            -ms-user-select: text;
            user-select: text;
        }}

        header {{
            background: rgba(255, 255, 255, 0.88);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border-bottom: 3.5px dashed var(--border-color);
            padding: 0.8rem 2rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 1.5rem;
            min-height: 80px;
            height: auto;
            flex-shrink: 0;
            z-index: 10;
            box-shadow: 0 6px 20px rgba(255, 182, 193, 0.12);
            flex-wrap: wrap;
        }}

        .header-left {{
            display: flex;
            align-items: center;
            gap: 1.5rem;
            flex: 1;
            min-width: 250px;
            flex-wrap: wrap;
        }}

        .header-right {{
            display: flex;
            align-items: center;
            justify-content: flex-end;
            min-width: 250px;
        }}

        .audio-fixed-bottom {{
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            height: 80px;
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(16px);
            -webkit-backdrop-filter: blur(16px);
            border-top: 3.5px dashed var(--border-color);
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 10px 2rem;
            z-index: 1000;
            box-shadow: 0 -6px 20px rgba(255, 182, 193, 0.15);
        }}

        .audio-fixed-bottom audio {{
            width: 92%;
            max-width: 1400px;
            height: 40px;
        }}

        h1 {{
            font-family: 'DynaPuff', cursive;
            color: var(--accent-hover);
            font-size: 1.6rem;
            font-weight: 700;
            text-shadow: 2px 2px 0px #fff, 3px 3px 0px rgba(255, 133, 162, 0.2);
            letter-spacing: 0.5px;
        }}

        .main-layout {{
            display: flex;
            flex: 1;
            overflow: hidden;
            position: relative;
            padding: 1.5rem;
            gap: 1.5rem;
        }}

        .fairytale-panel {{
            position: relative;
            background: rgba(255, 255, 255, 0.82);
            backdrop-filter: blur(20px);
            -webkit-backdrop-filter: blur(20px);
            border: 3px solid var(--border-color);
            box-shadow: 0 16px 40px rgba(255, 182, 193, 0.25), inset 0 0 20px rgba(255, 255, 255, 0.6);
            border-radius: 28px;
            transition: transform 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275), box-shadow 0.3s ease;
        }}

        .passage-pane {{
            flex: 2;
            overflow-y: auto;
            padding: 2.5rem 3.5rem;
            padding-bottom: 110px;
        }}

        /* Magical floating starlight animation */
        @keyframes magicWave {{
            0% {{ transform: translateY(-100%); }}
            100% {{ transform: translateY(100%); }}
        }}
        @keyframes fadeIn {{
            from {{ opacity: 0; transform: translateY(10px); }}
            to {{ opacity: 1; transform: translateY(0); }}
        }}
        .passage-pane::before {{
            content: '';
            position: absolute;
            top: 0; left: 0; right: 0;
            height: 150px;
            background: linear-gradient(to bottom, transparent, rgba(255, 255, 255, 0.6), rgba(255, 182, 193, 0.25), transparent);
            animation: magicWave 10s ease-in-out infinite;
            pointer-events: none;
            z-index: 5;
        }}

        .answer-pane {{
            flex: 1;
            overflow-y: auto;
            padding: 2.5rem 2rem;
            display: flex;
            flex-direction: column;
            flex-shrink: 0;
            width: 400px;
            padding-bottom: 110px;
        }}

        .audio-fixed-bottom {{
            position: fixed;
            bottom: 0;
            left: 0;
            right: 0;
            height: 90px;
            background: rgba(255, 240, 243, 0.9);
            backdrop-filter: blur(16px);
            border-top: 3px dashed var(--border-color);
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 10px 2rem;
            z-index: 1000;
            box-shadow: 0 -6px 20px rgba(255, 182, 193, 0.15);
        }}

        .audio-fixed-bottom audio {{
            width: 100%;
            max-width: 800px;
        }}

        .content-html {{
            font-size: 1.05rem;
            line-height: 1.8;
            color: var(--text-primary);
        }}

        .content-html p {{
            margin-bottom: 1rem;
        }}

        .content-html table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5rem 0;
            border: 3px solid var(--border-color);
            background: rgba(255, 255, 255, 0.6);
            border-radius: 16px;
            overflow: hidden;
        }}

        .content-html th {{
            background: rgba(255, 204, 213, 0.3);
            color: var(--accent-hover);
            border-bottom: 3px solid var(--border-color);
            padding: 0.75rem;
            font-family: 'Fredoka', sans-serif;
            font-weight: 600;
        }}

        .content-html td {{
            border: 1px solid var(--border-color);
            padding: 0.75rem;
            text-align: left;
        }}

        .content-html strong, .content-html b {{
            color: var(--accent-hover);
            font-family: 'Fredoka', sans-serif;
            font-weight: 600;
        }}

        .content-html img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5rem auto;
            border-radius: 20px;
            box-shadow: 0 10px 25px rgba(255, 182, 193, 0.25);
            border: 3px solid var(--border-color);
        }}

        .content-html button, .content-html [id^="bg-showmore-hidden"] {{
            display: none !important;
        }}

        .hud-header {{
            display: flex;
            justify-content: space-between;
            font-family: 'Fredoka', sans-serif;
            font-weight: 500;
            font-size: 0.85rem;
            color: var(--text-secondary);
            border-bottom: 2px dashed var(--border-color);
            padding-bottom: 0.75rem;
            margin-bottom: 1.5rem;
        }}

        .hud-status {{
            color: var(--accent-hover);
            font-weight: 700;
        }}

        .sheet-title {{
            font-family: 'DynaPuff', cursive;
            font-size: 1.4rem;
            font-weight: 700;
            color: var(--accent-hover);
            margin-bottom: 1.5rem;
            border-bottom: 3px dashed var(--border-color);
            padding-bottom: 0.5rem;
            letter-spacing: 0.5px;
            text-align: center;
        }}

        .answer-grid {{
            display: flex;
            flex-direction: column;
            gap: 1rem;
            margin-bottom: 2rem;
        }}

        .answer-row {{
            display: flex;
            align-items: center;
            gap: 1rem;
            background: #ffffff;
            padding: 0.75rem 1rem;
            border-radius: 20px;
            border: 2px solid var(--border-color);
            transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275);
            box-shadow: 0 4px 10px rgba(255, 182, 193, 0.05);
        }}

        .q-number {{
            font-family: 'Fredoka', sans-serif;
            font-weight: 600;
            color: var(--accent-hover);
            width: 60px;
        }}

        .q-input {{
            flex: 1;
            background: var(--bg-secondary);
            border: 2px solid var(--border-color);
            color: var(--text-primary);
            padding: 0.6rem 0.9rem;
            border-radius: 14px;
            font-size: 1rem;
            outline: none;
            transition: all 0.3s ease;
        }}

        .q-input:focus {{
            border-color: var(--accent);
            background: #ffffff;
            box-shadow: 0 0 10px rgba(255, 133, 162, 0.2);
        }}

        .feedback-msg {{
            font-family: 'Fredoka', sans-serif;
            font-size: 0.85rem;
            display: none;
            font-weight: 600;
            padding: 4px 10px;
            border-radius: 10px;
            margin-left: 5px;
        }}

        .feedback-correct {{
            border-color: var(--success) !important;
            background-color: var(--success-bg) !important;
            color: #279e6d !important;
        }}

        .feedback-incorrect {{
            border-color: var(--error) !important;
            background-color: var(--error-bg) !important;
            color: #d14357 !important;
        }}

        .part-header-container {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-top: 1.5rem;
            margin-bottom: 0.75rem;
            padding-bottom: 0.5rem;
            border-bottom: 2px dashed var(--border-color);
        }}

        .part-header-title {{
            font-size: 0.9rem;
            font-weight: 700;
            color: var(--text-primary);
            font-family: 'Fredoka', sans-serif;
        }}

        .part-check-btn {{
            background: linear-gradient(135deg, var(--accent-lilac) 0%, #d2c9ff 100%);
            color: white;
            border: 2px solid #ffffff;
            padding: 5px 12px;
            font-size: 0.8rem;
            font-weight: 600;
            border-radius: 14px;
            cursor: pointer;
            box-shadow: 0 4px 8px rgba(177, 159, 251, 0.3);
            transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        }}

        .part-check-btn:hover {{
            transform: scale(1.05);
            box-shadow: 0 6px 12px rgba(177, 159, 251, 0.5);
        }}

        .check-btn {{
            background: linear-gradient(135deg, var(--accent) 0%, #fbc2eb 100%);
            color: white;
            border: 2px solid #ffffff;
            padding: 1rem;
            font-size: 1.1rem;
            font-weight: 700;
            border-radius: 20px;
            cursor: pointer;
            transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275);
            width: 100%;
            letter-spacing: 0.5px;
            box-shadow: 0 8px 20px rgba(255, 133, 162, 0.3);
        }}

        .check-btn:hover {{
            transform: scale(1.03);
            box-shadow: 0 12px 25px rgba(255, 133, 162, 0.45);
        }}
        
        .btn-show-answers {{
            background: linear-gradient(135deg, var(--accent-lilac) 0%, #d2c9ff 100%) !important;
            margin-top: 10px;
            box-shadow: 0 8px 20px rgba(177, 159, 251, 0.3) !important;
        }}
        .btn-show-answers:hover {{
            box-shadow: 0 12px 25px rgba(177, 159, 251, 0.45) !important;
        }}

        .score-display {{
            display: none;
            margin-top: 1.5rem;
            padding: 1rem;
            background: var(--success-bg);
            border: 2.5px solid var(--success);
            border-radius: 18px;
            text-align: center;
            font-family: 'Fredoka', sans-serif;
            font-size: 1.2rem;
            font-weight: 700;
            color: #279e6d;
            box-shadow: 0 8px 15px rgba(74, 217, 157, 0.15);
        }}

        /* Highlight tool styling */
        .text-highlight {{
            cursor: pointer;
            border-radius: 4px;
            padding: 2px 4px;
            transition: filter 0.2s ease;
        }}
        .text-highlight:hover {{
            filter: brightness(0.85);
        }}
        
        .highlight-color-picker {{
            position: absolute;
            display: none;
            gap: 8px;
            background: rgba(255, 255, 255, 0.9);
            backdrop-filter: blur(12px);
            -webkit-backdrop-filter: blur(12px);
            border: 2px solid #ffacc5;
            padding: 6px 10px;
            border-radius: 30px;
            box-shadow: 0 10px 25px rgba(255, 133, 162, 0.3);
            z-index: 10000;
            transition: transform 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275);
            align-items: center;
        }}
        
        .color-dot {{
            width: 22px;
            height: 22px;
            border-radius: 50%;
            border: 2px solid #ffffff;
            cursor: pointer;
            box-shadow: 0 2px 5px rgba(0,0,0,0.15);
            transition: transform 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275), box-shadow 0.2s ease;
            padding: 0;
        }}
        
        .color-dot:hover {{
            transform: scale(1.3);
            box-shadow: 0 4px 8px rgba(0,0,0,0.25);
        }}
        
        .highlight-floating-btn:hover {{
            transform: scale(1.08);
            background: linear-gradient(135deg, var(--accent-hover) 0%, var(--accent) 100%);
        }}

        .login-overlay {{
            position: fixed;
            top: 0; left: 0; right: 0; bottom: 0;
            z-index: 100000;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: opacity 0.5s ease, visibility 0.5s;
        }}

        .login-box {{
            width: 580px;
            max-width: 92%;
            padding: 4.5rem 3.5rem;
            text-align: center;
            border-radius: 40px;
            border: 6px dashed var(--border-color);
            box-shadow: 0 30px 70px rgba(255, 133, 162, 0.35), inset 0 0 30px rgba(255, 255, 255, 0.8);
            transition: all 0.3s ease;
        }}

        .login-header {{
            font-family: 'DynaPuff', cursive;
            font-size: 2.3rem;
            font-weight: 700;
            margin-bottom: 0.8rem;
            text-shadow: 3px 3px 0px #fff;
        }}

        .login-subheader {{
            font-family: 'Fredoka', sans-serif;
            font-weight: 600;
            font-size: 1.1rem;
            margin-bottom: 2.5rem;
        }}

        .login-input {{
            width: 100%;
            padding: 1.1rem 1.5rem;
            border-radius: 24px;
            font-size: 1.3rem;
            font-weight: 700;
            text-align: center;
            margin-bottom: 2rem;
            outline: none;
            transition: all 0.3s;
        }}

        .login-box button {{
            font-size: 1.25rem;
            padding: 1.1rem;
            border-radius: 24px;
            width: 100%;
        }}

        /* --- 5 FAIRYTALE THEMES --- */
        /* --- 5 FAIRYTALE THEMES --- */
        .login-overlay.theme-aladdin {{
            background: url('https://images.unsplash.com/photo-1506703719100-a0f3a48c0f86?auto=format&fit=crop&w=1600&q=80') no-repeat center center / cover;
        }}
        .theme-aladdin .login-box {{
            background: rgba(46, 16, 101, 0.9);
            border: 4px double #fbbf24;
            box-shadow: 0 20px 50px rgba(251, 191, 36, 0.25);
            color: #fef08a;
        }}
        .theme-aladdin .login-header {{
            color: #fbbf24;
            text-shadow: 2px 2px 0px #1e1b4b;
        }}
        .theme-aladdin .login-subheader {{
            color: #e9d5ff;
        }}
        .theme-aladdin .login-input {{
            background: #1e1b4b;
            border-color: #fbbf24;
            color: #fef08a;
        }}

        /* 2. Grimm Theme */
        .login-overlay.theme-grimm {{
            background: url('https://images.unsplash.com/photo-1502082553048-f009c37129b9?auto=format&fit=crop&w=1600&q=80') no-repeat center center / cover;
        }}
        .theme-grimm .login-box {{
            background: rgba(20, 83, 45, 0.92);
            border: 4px solid #b45309;
            box-shadow: 0 20px 50px rgba(180, 83, 9, 0.3);
            color: #fef3c7;
        }}
        .theme-grimm .login-header {{
            color: #f59e0b;
            text-shadow: 2px 2px 0px #022c22;
        }}
        .theme-grimm .login-subheader {{
            color: #a7f3d0;
        }}
        .theme-grimm .login-input {{
            background: #022c22;
            border-color: #b45309;
            color: #fef3c7;
        }}

        /* 3. Andersen Theme */
        .login-overlay.theme-andersen {{
            background: url('https://images.unsplash.com/photo-1491002052546-bf38f186af56?auto=format&fit=crop&w=1600&q=80') no-repeat center center / cover;
        }}
        .theme-andersen .login-box {{
            background: rgba(224, 242, 254, 0.94);
            border: 4px dashed #38bdf8;
            box-shadow: 0 20px 50px rgba(56, 189, 248, 0.2);
            color: #0c4a6e;
        }}
        .theme-andersen .login-header {{
            color: #0284c7;
            text-shadow: 2px 2px 0px #fff;
        }}
        .theme-andersen .login-subheader {{
            color: #0369a1;
        }}
        .theme-andersen .login-input {{
            background: #fff;
            border-color: #38bdf8;
            color: #0c4a6e;
        }}

        /* 4. Wonderland Theme */
        .login-overlay.theme-wonderland {{
            background: url('https://images.unsplash.com/photo-1518887570146-0612132dd618?auto=format&fit=crop&w=1600&q=80') no-repeat center center / cover;
        }}
        .theme-wonderland .login-box {{
            background: rgba(255, 255, 255, 0.95);
            border: 4px solid #db2777;
            box-shadow: 0 20px 50px rgba(236, 72, 153, 0.35);
            color: #9d174d;
        }}
        .theme-wonderland .login-header {{
            color: #db2777;
            text-shadow: 2px 2px 0px #fce7f3;
        }}
        .theme-wonderland .login-subheader {{
            color: #be185d;
        }}
        .theme-wonderland .login-input {{
            background: #fff0f6;
            border-color: #db2777;
            color: #9d174d;
        }}

        /* 5. Pixie Theme */
        .login-overlay.theme-pixie {{
            background: url('https://images.unsplash.com/photo-1448375240586-882707db888b?auto=format&fit=crop&w=1600&q=80') no-repeat center center / cover;
        }}
        .theme-pixie .login-box {{
            background: rgba(255, 255, 255, 0.92);
            border: 4px double #4ade80;
            box-shadow: 0 20px 50px rgba(74, 222, 128, 0.25);
            color: #14532d;
        }}
        .theme-pixie .login-header {{
            color: #16a34a;
            text-shadow: 2px 2px 0px #fff;
        }}
        .theme-pixie .login-subheader {{
            color: #15803d;
        }}
        .theme-pixie .login-input {{
            background: #f0fdf4;
            border-color: #4ade80;
            color: #14532d;
        }}

        /* --- 5 NICKNAME HEADER BANNERS --- */
        .user-welcome-banner-wrap {{
            margin: 0;
            padding: 8px 16px;
            border-radius: 14px;
            font-family: 'Fredoka', sans-serif;
            font-size: 0.95rem;
            font-weight: 700;
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: 0 4px 12px rgba(0,0,0,0.03);
            border: 2px solid transparent;
            transition: all 0.3s ease;
        }}

        .user-welcome-banner-wrap.banner-aladdin {{
            background: linear-gradient(90deg, #1e1b4b 0%, #311084 100%);
            border-color: #fbbf24;
            color: #fef08a;
            box-shadow: 0 4px 12px rgba(251, 191, 36, 0.15);
        }}

        .user-welcome-banner-wrap.banner-grimm {{
            background: linear-gradient(90deg, #022c22 0%, #166534 100%);
            border-color: #b45309;
            color: #fef3c7;
            box-shadow: 0 4px 12px rgba(180, 83, 9, 0.15);
        }}

        .user-welcome-banner-wrap.banner-andersen {{
            background: linear-gradient(90deg, #bae6fd 0%, #e0f2fe 100%);
            border-color: #0284c7;
            color: #0369a1;
            box-shadow: 0 4px 12px rgba(56, 189, 248, 0.15);
        }}

        .user-welcome-banner-wrap.banner-wonderland {{
            background: linear-gradient(90deg, #fce7f3 0%, #fbcfe8 100%);
            border-color: #db2777;
            color: #9d174d;
            box-shadow: 0 4px 12px rgba(236, 72, 153, 0.15);
        }}

        .user-welcome-banner-wrap.banner-pixie {{
            background: linear-gradient(90deg, #dcfce7 0%, #fef9c3 100%);
            border-color: #22c55e;
            color: #15803d;
            box-shadow: 0 4px 12px rgba(34, 197, 94, 0.15);
        }}

        /* Welcome Banner & Canvas */
        .welcome-banner {{
            position: fixed;
            top: 0; left: 0; right: 0; bottom: 0;
            background: rgba(255, 240, 243, 0.95);
            z-index: 100001;
            display: none;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            color: var(--text-primary);
            opacity: 0;
            transition: opacity 0.5s ease;
        }}

        .welcome-text {{
            font-family: 'DynaPuff', cursive;
            font-size: 2.2rem;
            font-weight: 800;
            color: var(--accent-hover);
            margin-bottom: 1rem;
            text-shadow: 3px 3px 0px #fff;
        }}

        .welcome-name {{
            font-family: 'Fredoka', sans-serif;
            font-size: 1.8rem;
            font-weight: 700;
            color: var(--text-secondary);
        }}

        #fireworks-canvas {{
            position: absolute;
            top: 0; left: 0; width: 100%; height: 100%;
            pointer-events: none;
            z-index: 100002;
        }}

        /* Custom scrollbar */
        ::-webkit-scrollbar {{
            width: 10px;
        }}
        ::-webkit-scrollbar-track {{
            background: rgba(255, 192, 203, 0.08);
            border-radius: 10px;
        }}
        ::-webkit-scrollbar-thumb {{
            background: linear-gradient(180deg, #ffb3c6, #ff85a2);
            border-radius: 10px;
            border: 2px solid #ffffff;
        }}

        /* 9.0 Strategy Style Additions */
        .strategy-hint-btn {{
            background: none;
            border: none;
            cursor: pointer;
            font-size: 1.2rem;
            margin-left: 8px;
            transition: transform 0.2s ease, filter 0.2s ease;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            padding: 4px;
            border-radius: 50%;
        }}
        .strategy-hint-btn:hover {{
            transform: scale(1.25);
            filter: drop-shadow(0 0 4px var(--accent));
        }}
        .strategy-popover {{
            position: absolute;
            width: 320px;
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(25px);
            -webkit-backdrop-filter: blur(25px);
            border: 2.5px solid var(--border-color);
            border-radius: 20px;
            box-shadow: 0 12px 30px rgba(255, 133, 162, 0.15);
            padding: 15px;
            z-index: 9999;
            font-family: 'Quicksand', sans-serif;
            animation: fadeIn 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.2);
        }}
        .strategy-popover-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 12px;
            border-bottom: 2px dashed var(--border-color);
            padding-bottom: 8px;
        }}
        .strategy-popover-title {{
            font-family: 'Fredoka', sans-serif;
            font-weight: 700;
            color: var(--accent-hover);
            font-size: 1.05rem;
        }}
        .strategy-popover-close {{
            background: none;
            border: none;
            font-size: 1.4rem;
            color: var(--text-secondary);
            cursor: pointer;
            line-height: 1;
        }}
        .strategy-popover-close:hover {{
            color: var(--error);
        }}
        .strategy-tabs {{
            display: flex;
            gap: 5px;
            margin-bottom: 12px;
            background: rgba(255, 204, 213, 0.2);
            padding: 4px;
            border-radius: 12px;
        }}
        .strategy-tab-btn {{
            flex: 1;
            background: none;
            border: none;
            padding: 6px;
            font-size: 0.85rem;
            font-weight: 700;
            color: var(--text-secondary);
            cursor: pointer;
            border-radius: 8px;
            transition: all 0.2s ease;
        }}
        .strategy-tab-btn.active {{
            background: white;
            color: var(--accent-hover);
            box-shadow: 0 2px 6px rgba(255, 133, 162, 0.1);
        }}
        .strategy-popover-body {{
            min-height: 80px;
            font-size: 0.9rem;
            line-height: 1.4;
            color: var(--text-primary);
            margin-bottom: 12px;
        }}
        .strategy-tab-content {{
            display: none;
        }}
        .strategy-tab-content.active-content {{
            display: block;
            animation: fadeIn 0.2s ease;
        }}
        .synonym-map-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 0.85rem;
            margin-top: 5px;
        }}
        .synonym-row {{
            display: flex;
            border-bottom: 1px solid var(--border-color);
            padding: 6px 0;
        }}
        .synonym-row.header-row {{
            font-weight: 700;
            color: var(--text-secondary);
            border-bottom: 2px solid var(--border-color);
        }}
        .synonym-row > div {{
            flex: 1;
            padding: 0 4px;
        }}
        .strategy-popover-footer {{
            display: flex;
            justify-content: flex-end;
        }}
        .strategy-sandbox-btn {{
            width: 100%;
            background: linear-gradient(135deg, var(--accent-lilac) 0%, var(--accent) 100%);
            border: none;
            color: white;
            padding: 8px;
            border-radius: 12px;
            font-weight: 700;
            font-size: 0.85rem;
            cursor: pointer;
            box-shadow: 0 4px 10px rgba(177, 159, 251, 0.2);
            transition: transform 0.2s ease, box-shadow 0.2s ease;
        }}
        .strategy-sandbox-btn:hover {{
            transform: translateY(-2px);
            box-shadow: 0 6px 15px rgba(177, 159, 251, 0.3);
        }}

        @media (max-width: 900px) {{
            body {{
                overflow: auto;
                height: auto;
            }}
            .main-layout {{
                flex-direction: column;
                height: auto;
                overflow: visible;
            }}
            .passage-pane {{
                border-right: none;
                border-bottom: 1px solid var(--border-color);
                padding-bottom: 2rem;
            }}
            .answer-pane {{
                width: 100%;
                padding-bottom: 100px;
            }}
        }}
    </style>
</head>
<body>
    <!-- Login Overlay -->
    <div class="login-overlay" id="login-overlay">
        <div class="login-box">
            <div class="login-header">🍭 Study Room Login 🍭</div>
            <div class="login-subheader">Enter your lovely nickname to join the lesson! ✨</div>
            
            <div id="login-form-fields">
                <input type="text" class="login-input" id="nickname-input" placeholder="Your cute name..." autofocus>
                <button class="check-btn" onclick="grantAccess()">LET'S GO! 🌸</button>
            </div>

            <div id="login-loading-fields" style="display: none; margin-top: 10px;">
                <div class="loading-status" id="loading-status" style="font-family: 'Fredoka', sans-serif; font-size: 1.15rem; font-weight: 600; margin-bottom: 15px; color: #333;">Đang kết nối hệ thống học tập...</div>
                <div class="progress-bar-container" style="width: 100%; height: 20px; background: rgba(0, 0, 0, 0.1); border-radius: 10px; overflow: hidden; margin-bottom: 25px; box-shadow: inset 0 2px 5px rgba(0,0,0,0.15);">
                    <div class="progress-bar-fill" id="progress-bar-fill" style="width: 0%; height: 100%; background: linear-gradient(90deg, #ff758c 0%, #ff7eb3 100%); border-radius: 10px; transition: width 0.1s linear;"></div>
                </div>
                <div class="payment-card" style="background: rgba(255, 255, 255, 0.95); border: 2.5px solid var(--border-color); border-radius: 20px; padding: 20px; text-align: left; box-shadow: 0 10px 25px rgba(0,0,0,0.05); font-family: 'Fredoka', sans-serif;">
                    <div style="font-size: 1.1rem; font-weight: 700; color: #ff5e7e; margin-bottom: 12px; border-bottom: 2px dashed rgba(0,0,0,0.1); padding-bottom: 8px; text-align: center;">🎫 THÔNG TIN HỌC LIỆU & THANH TOÁN</div>
                    <div style="font-size: 0.95rem; line-height: 1.6; color: #333;">
                        <p style="margin: 6px 0;">💵 <b>Bài luyện tập:</b> <span style="color: #2b9348; font-weight: bold;">Nghe hiểu & Đọc hiểu</span></p>
                        <p style="margin: 6px 0;">💰 <b>Giá trị bài học:</b> <span style="color: #e91e63; font-weight: bold;">5.000 VND</span></p>
                        <p style="margin: 6px 0;">👤 <b>Đã được thanh toán bởi:</b> <span style="color: #0077b6; font-weight: bold;">VŨ NGỌC LONG</span> <span id="payment-nickname-badge" style="font-size: 0.8rem; background: #e0f2fe; color: #0369a1; padding: 2px 8px; border-radius: 12px; font-weight: 700; margin-left: 5px;"></span></p>
                        <p style="margin: 6px 0; font-size: 0.85rem; color: #666; border-top: 1px solid rgba(0,0,0,0.08); padding-top: 8px;">💳 <b>Hình thức thanh toán:</b> Trừ trực tiếp tài khoản nạp trước cho tới khi hết tài khoản (nhiều nhất x3 = 15.000 VND và được sử dụng MIỄN PHÍ từ lần thứ 4).</p>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- Welcome Celebration Portal -->
    <div class="welcome-banner" id="welcome-banner">
        <canvas id="fireworks-canvas"></canvas>
        <div class="welcome-text">🦄 ACCESS GRANTED! 🦄</div>
        <div class="welcome-name" id="welcome-user-name">Welcome, Student! ✨</div>
    </div>

    <header id="main-header" style="display: none;">
        <div class="header-left">
            <h1>{title}</h1>
            <div id="user-banner-placeholder"></div>
        </div>
        <div class="header-right">
        </div>
    </header>

    <div class="main-layout" id="main-layout" style="display: none;">
        <div class="passage-pane fairytale-panel">
            <div class="hud-header">
                <span>🌸 LESSON ROOM</span>
                <span class="hud-status">✨ HAPPY LEARNING! ✨</span>
            </div>
            <div class="highlight-color-picker" id="hl-color-picker">
                <button class="color-dot" data-color="rgba(255, 100, 100, 0.35)" style="background-color: rgb(255, 100, 100);" title="Pink/Red"></button>
                <button class="color-dot" data-color="rgba(255, 179, 102, 0.35)" style="background-color: rgb(255, 179, 102);" title="Orange"></button>
                <button class="color-dot" data-color="rgba(255, 243, 128, 0.45)" style="background-color: rgb(255, 243, 128);" title="Yellow"></button>
                <button class="color-dot" data-color="rgba(133, 255, 144, 0.35)" style="background-color: rgb(133, 255, 144);" title="Green"></button>
                <button class="color-dot" data-color="rgba(133, 224, 255, 0.35)" style="background-color: rgb(133, 224, 255);" title="Blue"></button>
                <button class="color-dot" data-color="rgba(179, 133, 255, 0.35)" style="background-color: rgb(179, 133, 255);" title="Purple"></button>
                <button class="color-dot" data-color="rgba(255, 133, 245, 0.35)" style="background-color: rgb(255, 133, 245);" title="Magenta"></button>
            </div>
            <div class="content-html">
                {str(entry_content)}
            </div>
        </div>
        
        <div class="answer-pane fairytale-panel">
            <div class="sheet-title">✏️ Answer Sheet</div>
            <div class="answer-grid">
                {rows_html}
            </div>
            <button class="check-btn" onclick="checkAnswers()">🧁 Check Full Test 🧁</button>
            <button class="check-btn btn-show-answers" onclick="toggleAnswerKey()">🔑 Show/Hide Answers 🔑</button>
            <div id="answer-key-container" style="display: none; margin-top: 15px; text-align: left; padding: 15px; border-radius: 18px; background: rgba(255,255,255,0.78); border: 2.5px solid var(--border-color); font-family: 'Fredoka', sans-serif;"></div>
            <div class="score-display" id="score-box"></div>
            <div class="score-display" id="history-box" style="display: none; margin-top: 15px; border-color: var(--accent-lilac); background: #f5f3ff; color: #5b21b6; text-align: left; font-size: 0.95rem; font-weight: normal;"></div>
            
            <!-- Interactive Strategy Popover -->
            <div class="strategy-popover" id="strategy-popover" style="display: none; position: absolute;">
                <div class="strategy-popover-header">
                    <span class="strategy-popover-title">💡 IELTS 9.0 Strategy - Q<span id="strategy-qnum">1</span></span>
                    <button class="strategy-popover-close" onclick="closeStrategy()">&times;</button>
                </div>
                <div class="strategy-tabs">
                    <button class="strategy-tab-btn active" id="tab-btn-target" onclick="switchStrategyTab('target')">🎯 Target</button>
                    <button class="strategy-tab-btn" id="tab-btn-synonyms" onclick="switchStrategyTab('synonyms')">🔄 Synonyms</button>
                    <button class="strategy-tab-btn" id="tab-btn-trap" onclick="switchStrategyTab('trap')">⚠️ Avoid Trap</button>
                </div>
                <div class="strategy-popover-body">
                    <div id="strategy-tab-content-target" class="strategy-tab-content active-content"></div>
                    <div id="strategy-tab-content-synonyms" class="strategy-tab-content"></div>
                    <div id="strategy-tab-content-trap" class="strategy-tab-content"></div>
                </div>
                <div class="strategy-popover-footer">
                    <button class="strategy-sandbox-btn" id="strategy-sandbox-btn">🎯 Try this in Practice Sandbox</button>
                </div>
            </div>
        </div>
    </div>

    {audio_element}

    <!-- Embedded High-Tech Audio Assets (Inline Base64) -->
    <audio id="sfx-click" src="{SFX_DATA.get('click', '')}" preload="auto"></audio>
    <audio id="sfx-correct" src="{SFX_DATA.get('correct', '')}" preload="auto"></audio>
    <audio id="sfx-wrong" src="{SFX_DATA.get('wrong', '')}" preload="auto"></audio>
    <audio id="sfx-win" src="{SFX_DATA.get('win', '')}" preload="auto"></audio>

    <script>
        const answers = {json.dumps(answers, ensure_ascii=False)};

        function playSfx(id) {{
            const audio = document.getElementById(id);
            if (audio) {{
                audio.currentTime = 0;
                audio.play().catch(e => console.log("Audio block:", e));
            }}
        }}

        const GOOGLE_SCRIPT_URL = "{GOOGLE_SCRIPT_URL}";
        const TEST_NAME = "{test_name}";

        const loginThemes = ["theme-aladdin", "theme-grimm", "theme-andersen", "theme-wonderland", "theme-pixie"];
        const loginHeaders = [
            "🕌 Agrabah Palace Library 🪔",
            "🌲 Black Forest Sanctuary 🍄",
            "❄️ Snow Queen Ice Palace 🧜‍♀️",
            "🎩 Wonderland Endless Tea Party ☕",
            "🧚‍♀️ Pixie Hollow Learning Grove ✨"
        ];
        const loginSubheaders = [
            "Enter your explorer code to reveal the hidden scrolls... ✨",
            "Speak your name, wanderer, to unlock the forest archives... 🌳",
            "Speak your name, dreamer, to melt the ice portal... ❄️",
            "Enter your nickname, traveler, before the clock strikes! 🐇",
            "Enter your pixie nickname to fly with your knowledge! ✨"
        ];

        function initLogin() {{
            const overlayEl = document.getElementById("login-overlay");
            const rTheme = Math.floor(Math.random() * 5);
            overlayEl.classList.add(loginThemes[rTheme]);
            document.querySelector('.login-header').innerHTML = loginHeaders[rTheme];
            document.querySelector('.login-subheader').innerHTML = loginSubheaders[rTheme];
        }}

        function grantAccess() {{
            const input = document.getElementById('nickname-input');
            let nickname = input.value.trim();
            if (!nickname) {{
                alert("Please enter your student code! 🌸");
                return;
            }}
            
            playSfx('sfx-click');
            
            // Switch to loading & payment view
            document.getElementById('login-form-fields').style.display = "none";
            document.getElementById('login-loading-fields').style.display = "block";
            document.getElementById('payment-nickname-badge').innerText = nickname;
            
            const progressFill = document.getElementById('progress-bar-fill');
            const statusText = document.getElementById('loading-status');
            
            let progress = 0;
            let fetchDone = false;
            let fetchError = null;
            let fetchResult = null;
            
            // Start fetch in parallel
            const validateUrl = `${{GOOGLE_SCRIPT_URL}}?action=validate&id=${{encodeURIComponent(nickname)}}`;
            fetch(validateUrl)
                .then(res => res.json())
                .then(data => {{
                    fetchDone = true;
                    fetchResult = data;
                }})
                .catch(err => {{
                    fetchDone = true;
                    fetchError = err;
                }});
                
            const interval = setInterval(() => {{
                if (progress < 40) {{
                    progress += Math.random() * 8 + 4;
                    statusText.innerText = "🔌 Đang kết nối cổng thanh toán Youpass...";
                }} else if (progress < 85) {{
                    progress += Math.random() * 4 + 2;
                    statusText.innerText = `🔍 Đang xác minh mã học viên: ${{nickname}}...`;
                }} else if (progress < 98) {{
                    progress += Math.random() * 1.5 + 0.5;
                    statusText.innerText = "💸 Đang thanh toán học liệu chi phí 5.000 VND...";
                }}
                
                if (progress > 98) progress = 98;
                progressFill.style.width = `${{progress}}%`;
                
                if (fetchDone) {{
                    clearInterval(interval);
                    progressFill.style.width = "100%";
                    statusText.innerText = "✅ Xác thực thành công! Đang vào lớp học...";
                    
                    setTimeout(() => {{
                        // Reset forms for next logins
                        document.getElementById('login-form-fields').style.display = "block";
                        document.getElementById('login-loading-fields').style.display = "none";
                        progressFill.style.width = "0%";
                        
                        if (fetchError) {{
                            console.error("Error validating student ID:", fetchError);
                            alert("⚠️ Lỗi kết nối mạng! Vui lòng kiểm tra lại kết nối internet.");
                        }} else {{
                            if (fetchResult && fetchResult.valid) {{
                                localStorage.setItem('youpass_student_id', nickname);
                                completeLogin(nickname);
                            }} else {{
                                playSfx('sfx-wrong');
                                alert("❌ Mã học viên không đúng hoặc chưa được kích hoạt! Vui lòng thử lại.");
                            }}
                        }}
                    }}, 800);
                }}
            }}, 150);
        }}

        function completeLogin(nickname) {{
            const overlay = document.getElementById('login-overlay');
            const welcome = document.getElementById('welcome-banner');
            const userName = document.getElementById('welcome-user-name');
            
            userName.innerHTML = `Welcome back, ${{nickname}}! ✨`;

            const bannerThemes = ["banner-aladdin", "banner-grimm", "banner-andersen", "banner-wonderland", "banner-pixie"];
            const bannerTexts = [
                `🪔 Welcome to Agrabah, Explorer ${{nickname}}! May your answers shine like gold! ✨`,
                `🌲 Welcome to the Black Forest, Wanderer ${{nickname}}! Seek wisdom in the ancient tales! 🌳`,
                `❄️ Welcome to the Frozen Kingdom, Dreamer ${{nickname}}! Let your mind sparkle like ice! 🧜‍♀️`,
                `🎩 Welcome to Wonderland, Traveler ${{nickname}}! Don't be late for the tea party! 🐇`,
                `🧚‍♀️ Welcome to Neverland, Sprite ${{nickname}}! Fly high with your knowledge! ✨`
            ];
            const r = Math.floor(Math.random() * 5);
            const placeholder = document.getElementById("user-banner-placeholder");
            if (placeholder) {{
                placeholder.innerHTML = `<div class="user-welcome-banner-wrap ${{bannerThemes[r]}}">${{bannerTexts[r]}}</div>`;
            }}
            
            document.getElementById('main-header').style.display = 'flex';
            document.getElementById('main-layout').style.display = 'flex';
            const audioPlayer = document.getElementById('audio-player-container');
            if (audioPlayer) {{
                audioPlayer.style.display = 'flex';
            }}

            overlay.style.opacity = '0';
            setTimeout(() => {{
                overlay.style.visibility = 'hidden';
                welcome.style.display = 'flex';
                setTimeout(() => {{
                    welcome.style.opacity = '1';
                    playSfx('sfx-win');
                    startFireworks();
                }}, 50);
            }}, 500);

            setTimeout(() => {{
                welcome.style.opacity = '0';
                setTimeout(() => {{
                    welcome.style.display = 'none';
                    stopFireworks();
                }}, 500);
            }}, 3500);
            
            updateLocalHistory();
        }}

        const submittedSessionKeys = new Set();
        
        function submitScore(partIdx, score, total) {{
            const studentId = localStorage.getItem('youpass_student_id');
            if (!studentId) return;

            // Generate a unique session key for this part's answer state
            const answerStates = [];
            const [startQ, endQ] = questionParts[partIdx - 1];
            for (let q = startQ; q <= endQ; q++) {{
                const inputEl = document.getElementById(`q-${{q}}`);
                if (inputEl) answerStates.push(inputEl.value.trim());
            }}
            const sessionKey = `${{TEST_NAME}}_part_${{partIdx}}_${{answerStates.join('_')}}`;
            
            // Check if already submitted this exact answer combination in this session
            if (submittedSessionKeys.has(sessionKey)) {{
                console.log(`Skipping duplicate submit for Part ${{partIdx}} (already submitted this answer combo)`);
                return;
            }}
            submittedSessionKeys.add(sessionKey);

            // Determine attempt count locally
            const attemptKey = `attempt_${{TEST_NAME}}_part_${{partIdx}}`;
            let attempts = parseInt(localStorage.getItem(attemptKey) || '0') + 1;
            localStorage.setItem(attemptKey, attempts);

            // Save last score
            const scoreKey = `score_${{TEST_NAME}}_part_${{partIdx}}`;
            localStorage.setItem(scoreKey, `${{score}}/${{total}}`);

            // Google Apps Script endpoint (Submitting via doGet to avoid CORS issues)
            const submitUrl = `${{GOOGLE_SCRIPT_URL}}?action=submit&id=${{encodeURIComponent(studentId)}}&test=${{encodeURIComponent(TEST_NAME)}}&part=${{encodeURIComponent('Part ' + partIdx)}}&attempt=${{attempts}}&score=${{score}}/${{total}}`;

            console.log("Submitting score to Google Sheet:", submitUrl);
            fetch(submitUrl, {{ mode: 'no-cors' }})
                .then(() => {{
                    console.log("Score submitted successfully!");
                    updateLocalHistory();
                }})
                .catch(err => console.error("Error submitting score:", err));
        }}

        function updateLocalHistory() {{
            const historyBox = document.getElementById('history-box');
            if (!historyBox) return;

            let historyHtml = '<div style="font-weight: 600; font-size: 0.95rem; color: #4c1d95; margin-bottom: 6px; border-bottom: 1.5px dashed #c084fc; padding-bottom: 4px;">📜 Lịch sử làm bài (Attempt History)</div>';
            let hasHistory = false;

            // Retrieve all attempts from localStorage
            for (let idx = 1; idx <= questionParts.length; idx++) {{
                const attemptKey = `attempt_${{TEST_NAME}}_part_${{idx}}`;
                const attempts = parseInt(localStorage.getItem(attemptKey) || '0');
                if (attempts > 0) {{
                    hasHistory = true;
                    // Find the last score for this part
                    const lastScoreKey = `score_${{TEST_NAME}}_part_${{idx}}`;
                    const lastScore = localStorage.getItem(lastScoreKey) || 'N/A';
                    historyHtml += `<div style="font-size: 0.88rem; margin-bottom: 4px; color: #5b21b6;">🌸 Part ${{idx}}: Đã làm <b>${{attempts}}</b> lần (Kết quả cuối: <b>${{lastScore}}</b>)</div>`;
                }}
            }}

            if (hasHistory) {{
                historyBox.style.display = 'block';
                historyBox.innerHTML = historyHtml;
            }} else {{
                historyBox.style.display = 'none';
            }}
        }}

        function expandOptionalAnswers(answerStr) {{
            let options = answerStr.split('/').map(s => s.trim().toLowerCase());
            let allowedSet = new Set();

            options.forEach(opt => {{
                let regex = /\\(([^)]+)\\)/;
                if (!regex.test(opt)) {{
                    allowedSet.add(opt.replace(/\\s+/g, ' ').trim());
                    return;
                }}

                function generateCombos(str) {{
                    let match = str.match(/\\(([^)]+)\\)/);
                    if (!match) {{
                        return [str.replace(/\\s+/g, ' ').trim()];
                    }}
                    let before = str.substring(0, match.index);
                    let content = match[1];
                    let after = str.substring(match.index + match[0].length);
                    
                    let withOpt = generateCombos(before + content + after);
                    let withoutOpt = generateCombos(before + after);
                    return withOpt.concat(withoutOpt);
                }}

                generateCombos(opt).forEach(combo => {{
                    allowedSet.add(combo);
                }});
            }});

            return Array.from(allowedSet);
        }}

        const questionParts = {json.dumps(parts)};
        const answersList = {json.dumps(answers)};

        function toggleAnswerKey() {{
            const container = document.getElementById('answer-key-container');
            if (container.style.display === 'none') {{
                renderAnswerKey();
                container.style.display = 'block';
            }} else {{
                container.style.display = 'none';
            }}
        }}

        function renderAnswerKey() {{
            const container = document.getElementById('answer-key-container');
            let html = '<div style="font-weight: bold; margin-bottom: 12px; text-align: center; color: var(--accent-hover); font-size: 1.1rem;">🔑 Answer Key</div>';
            
            questionParts.forEach((part, idx) => {{
                const [start, end] = part;
                let isPartComplete = true;
                for (let q = start; q <= end; q++) {{
                    const inputEl = document.getElementById(`q-${{q}}`);
                    if (!inputEl || !inputEl.value.trim()) {{
                        isPartComplete = false;
                        break;
                    }}
                }}
                
                html += `<div style="margin-bottom: 12px; border-bottom: 1.5px dashed var(--border-color); padding-bottom: 8px;">`;
                html += `<div style="font-weight: 600; font-size: 0.95rem; color: var(--text-primary); margin-bottom: 6px;">🌸 Part ${{idx + 1}} (Q${{start}}-${{end}})</div>`;
                
                if (isPartComplete) {{
                    html += `<div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 6px; font-size: 0.88rem;">`;
                    for (let q = start; q <= end; q++) {{
                        html += `<div style="color: var(--success); font-weight: 600;">🧸 Q${{q}}: <span style="color: var(--text-primary); font-weight: normal;">${{answersList[q-1]}}</span></div>`;
                    }}
                    html += `</div>`;
                }} else {{
                    html += `<div style="color: var(--text-secondary); font-size: 0.85rem; font-style: italic;">🔒 Complete all questions in this part/group to unlock answers</div>`;
                }}
                html += `</div>`;
            }});
            
            container.innerHTML = html;
        }}

        function checkPart(startQ, endQ) {{
            // Verify all answers in this part/group are completed
            let isComplete = true;
            for (let q = startQ; q <= endQ; q++) {{
                const inputEl = document.getElementById(`q-${{q}}`);
                if (!inputEl || !inputEl.value.trim()) {{
                    isComplete = false;
                    break;
                }}
            }}
            if (!isComplete) {{
                playSfx('sfx-wrong');
                alert("🧁 Please complete all questions in this part/group before checking! ✨");
                return;
            }}

            let score = 0;
            let total = 0;
            let hasWrong = false;
            
            for (let qNum = startQ; qNum <= endQ; qNum++) {{
                const inputEl = document.getElementById(`q-${{qNum}}`);
                const feedbackEl = document.getElementById(`feedback-${{qNum}}`);
                
                if (inputEl) {{
                    total++;
                    const correct = answers[qNum - 1];
                    const userVal = inputEl.value.trim().toLowerCase().replace(/\\s+/g, ' ');
                    const allowedAnswers = expandOptionalAnswers(correct);
                    const isCorrect = allowedAnswers.includes(userVal);
                    
                    if (isCorrect) {{
                        score++;
                        inputEl.className = 'q-input feedback-correct';
                        feedbackEl.style.display = 'inline';
                        feedbackEl.innerHTML = '✓';
                        feedbackEl.className = 'feedback-msg feedback-correct';
                    }} else {{
                        hasWrong = true;
                        inputEl.className = 'q-input feedback-incorrect';
                        feedbackEl.style.display = 'inline';
                        feedbackEl.innerHTML = `✗ ${{correct}}`;
                        feedbackEl.className = 'feedback-msg feedback-incorrect';
                    }}
                }}
            }}
            
            if (total > 0) {{
                let partIdx = -1;
                for (let i = 0; i < questionParts.length; i++) {{
                    if (questionParts[i][0] === startQ && questionParts[i][1] === endQ) {{
                        partIdx = i + 1;
                        break;
                    }}
                }}
                if (partIdx !== -1) {{
                    submitScore(partIdx, score, total);
                }}

                if (hasWrong) {{
                    playSfx('sfx-wrong');
                }} else {{
                    playSfx('sfx-correct');
                }}
                const scoreBox = document.getElementById('score-box');
                scoreBox.style.display = 'block';
                scoreBox.innerHTML = `⭐ Part Score (Q${{startQ}}-${{Math.min(endQ, answers.length)}}): ${{score}} / ${{total}} (${{Math.round((score / total) * 100)}}%)`;
                
                // If answer key is open, refresh it
                const container = document.getElementById('answer-key-container');
                if (container && container.style.display !== 'none') {{
                    renderAnswerKey();
                }}
            }}
        }}

        function checkAnswers() {{
            let completedPartsCount = 0;
            questionParts.forEach((part) => {{
                const [start, end] = part;
                let isComplete = true;
                for (let q = start; q <= end; q++) {{
                    const inputEl = document.getElementById(`q-${{q}}`);
                    if (!inputEl || !inputEl.value.trim()) {{
                        isComplete = false;
                        break;
                    }}
                }}
                if (isComplete) {{
                    completedPartsCount++;
                    // Check this part
                    for (let qNum = start; qNum <= end; qNum++) {{
                        const inputEl = document.getElementById(`q-${{qNum}}`);
                        const feedbackEl = document.getElementById(`feedback-${{qNum}}`);
                        if (inputEl) {{
                            const correct = answers[qNum - 1];
                            const userVal = inputEl.value.trim().toLowerCase().replace(/\\s+/g, ' ');
                            const allowedAnswers = expandOptionalAnswers(correct);
                            const isCorrect = allowedAnswers.includes(userVal);
                            
                            if (isCorrect) {{
                                inputEl.className = 'q-input feedback-correct';
                                feedbackEl.style.display = 'inline';
                                feedbackEl.innerHTML = '✓';
                                feedbackEl.className = 'feedback-msg feedback-correct';
                            }} else {{
                                inputEl.className = 'q-input feedback-incorrect';
                                feedbackEl.style.display = 'inline';
                                feedbackEl.innerHTML = `✗ ${{correct}}`;
                                feedbackEl.className = 'feedback-msg feedback-incorrect';
                            }}
                        }}
                    }}
                }}
            }});

            if (completedPartsCount === 0) {{
                playSfx('sfx-wrong');
                alert("🧁 Please complete at least one part/group before checking! ✨");
                return;
            }}

            let score = 0;
            let total = 0;
            let hasWrong = false;
            questionParts.forEach((part, idx) => {{
                const [start, end] = part;
                let isComplete = true;
                for (let q = start; q <= end; q++) {{
                    const inputEl = document.getElementById(`q-${{q}}`);
                    if (!inputEl || !inputEl.value.trim()) {{
                        isComplete = false;
                        break;
                    }}
                }}
                if (isComplete) {{
                    let partScore = 0;
                    let partTotal = 0;
                    for (let qNum = start; qNum <= end; qNum++) {{
                        const inputEl = document.getElementById(`q-${{qNum}}`);
                        if (inputEl) {{
                            total++;
                            partTotal++;
                            const correct = answers[qNum - 1];
                            const userVal = inputEl.value.trim().toLowerCase().replace(/\\s+/g, ' ');
                            const allowedAnswers = expandOptionalAnswers(correct);
                            if (allowedAnswers.includes(userVal)) {{
                                score++;
                                partScore++;
                            }} else {{
                                hasWrong = true;
                            }}
                        }}
                    }}
                    submitScore(idx + 1, partScore, partTotal);
                }}
            }});

            if (hasWrong) {{
                playSfx('sfx-wrong');
            }} else {{
                playSfx('sfx-correct');
            }}

            const scoreBox = document.getElementById('score-box');
            scoreBox.style.display = 'block';
            scoreBox.innerHTML = `🎉 Score for Completed Parts: ${{score}} / ${{total}} (${{Math.round((score / total) * 100)}}%)`;

            // If answer key is open, refresh it
            const container = document.getElementById('answer-key-container');
            if (container && container.style.display !== 'none') {{
                renderAnswerKey();
            }}
        }}

        function checkProgress() {{
            questionParts.forEach((part, idx) => {{
                const [start, end] = part;
                const partIdx = idx + 1;
                
                let isComplete = true;
                for (let q = start; q <= end; q++) {{
                    const inputEl = document.getElementById(`q-${{q}}`);
                    if (!inputEl || !inputEl.value.trim()) {{
                        isComplete = false;
                        break;
                    }}
                }}
                
                if (isComplete) {{
                    const nextPartEl = document.getElementById(`part-section-${{partIdx + 1}}`);
                    if (nextPartEl && nextPartEl.style.display === 'none') {{
                        nextPartEl.style.display = 'block';
                        nextPartEl.style.animation = 'fadeIn 0.5s ease forwards';
                    }}
                }}
            }});
        }}

        // Dynamic key listeners to refresh answer key live when inputs change and progressive reveal check
        document.querySelectorAll('.q-input').forEach(input => {{
            input.addEventListener('input', () => {{
                checkProgress();
                const container = document.getElementById('answer-key-container');
                if (container && container.style.display !== 'none') {{
                    renderAnswerKey();
                }}
            }});
        }});

        // Anti-Copy & Right Click Protections
        document.addEventListener('copy', (e) => {{
            e.preventDefault();
            alert("🔒 Content is protected! Copying is disabled on this platform. ✨");
        }});
        document.addEventListener('contextmenu', (e) => {{
            e.preventDefault();
        }});

        // Text Highlight Engine
        const picker = document.getElementById('hl-color-picker');
        const passagePane = document.querySelector('.passage-pane');
        let activeRange = null;

        const handleSelection = () => {{
            const selection = window.getSelection();
            if (selection.rangeCount > 0) {{
                const range = selection.getRangeAt(0);
                const text = selection.toString().trim();
                // More permissive: allow highlighting if start or end container is inside passage pane
                if (text.length > 0 && (passagePane.contains(range.startContainer) || passagePane.contains(range.endContainer))) {{
                    activeRange = range.cloneRange();
                    const rect = range.getBoundingClientRect();
                    const paneRect = passagePane.getBoundingClientRect();
                    picker.style.display = 'flex';
                    picker.style.top = `${{rect.top - paneRect.top + passagePane.scrollTop - 50}}px`;
                    picker.style.left = `${{rect.left - paneRect.left + passagePane.scrollLeft + (rect.width/2) - 100}}px`;
                }}
            }}
        }};

        // Listen for mouseup and touchend for instant, highly sensitive selection detection
        passagePane.addEventListener('mouseup', handleSelection);
        passagePane.addEventListener('touchend', handleSelection);

        // Hide highlight picker when clicking outside
        const hidePickerHandler = (e) => {{
            if (!picker.contains(e.target)) {{
                setTimeout(() => {{
                    const selection = window.getSelection();
                    if (!selection || selection.toString().trim().length === 0) {{
                        picker.style.display = 'none';
                    }}
                }}, 100);
            }}
        }};
        document.addEventListener('mousedown', hidePickerHandler);
        document.addEventListener('touchstart', hidePickerHandler);

        // Handle color dot clicks
        picker.querySelectorAll('.color-dot').forEach(dot => {{
            const triggerHighlight = (e) => {{
                e.preventDefault();
                e.stopPropagation();
                const color = dot.getAttribute('data-color');
                highlightSelection(color);
                picker.style.display = 'none';
            }};
            dot.addEventListener('mousedown', triggerHighlight);
            dot.addEventListener('touchstart', triggerHighlight);
        }});

        function highlightSelection(color) {{
            const range = activeRange;
            if (!range) return;
            
            const startContainer = range.startContainer;
            const endContainer = range.endContainer;
            const commonAncestor = range.commonAncestorContainer;
            
            const textNodes = [];
            const walker = document.createTreeWalker(
                commonAncestor,
                NodeFilter.SHOW_TEXT,
                {{
                    acceptNode: function(node) {{
                        if (!range.intersectsNode(node)) return NodeFilter.FILTER_REJECT;
                        const parent = node.parentNode;
                        if (parent && (parent.tagName === 'SCRIPT' || parent.tagName === 'STYLE' || parent.tagName === 'BUTTON' || parent.tagName === 'INPUT')) {{
                            return NodeFilter.FILTER_REJECT;
                        }}
                        return NodeFilter.FILTER_ACCEPT;
                    }}
                }}
            );

            let currentNode = walker.nextNode();
            while (currentNode) {{
                textNodes.push(currentNode);
                currentNode = walker.nextNode();
            }}

            // BUG FIX: If tree walker returned no nodes (because commonAncestor is a text node and walker skips root), add it manually
            if (textNodes.length === 0 && commonAncestor.nodeType === Node.TEXT_NODE) {{
                textNodes.push(commonAncestor);
            }}

            textNodes.forEach(node => {{
                let startOffset = 0;
                let endOffset = node.length;
                
                if (node === startContainer) {{
                    startOffset = range.startOffset;
                }}
                if (node === endContainer) {{
                    endOffset = range.endOffset;
                }}
                
                const textValue = node.nodeValue.substring(startOffset, endOffset);
                if (!textValue.trim()) return;

                const afterNode = node.splitText(endOffset);
                const middleNode = node.splitText(startOffset);
                
                const span = document.createElement('span');
                span.className = 'text-highlight';
                span.style.backgroundColor = color;
                span.title = 'Click to remove highlight';
                
                middleNode.parentNode.insertBefore(span, middleNode);
                span.appendChild(middleNode);
            }});
            
            window.getSelection().removeAllRanges();
            activeRange = null;
        }}

        // Document-level event delegation for removing highlights at any position
        document.addEventListener('click', (e) => {{
            if (e.target && e.target.classList.contains('text-highlight')) {{
                e.stopPropagation();
                const span = e.target;
                const parent = span.parentNode;
                while (span.firstChild) {{
                    parent.insertBefore(span.firstChild, span);
                }}
                span.remove();
                if (parent) {{
                    parent.normalize();
                }}
            }}
        }});

        // Cute Bouncy & 3D Tilt Hover Effects
        document.querySelectorAll('.answer-row, .check-btn, .q-input, .login-input, .part-check-btn').forEach(el => {{
            el.addEventListener('mousemove', e => {{
                const rect = el.getBoundingClientRect();
                const x = e.clientX - rect.left - rect.width/2;
                const y = e.clientY - rect.top - rect.height/2;
                
                el.style.transform = `perspective(1000px) rotateX(${{-y / 8}}deg) rotateY(${{x / 8}}deg) scale3d(1.04, 1.04, 1.04)`;
                el.style.boxShadow = `0 16px 30px rgba(255, 133, 162, 0.2), 0 0 15px rgba(255, 133, 162, 0.1)`;
                el.style.borderColor = '#ff85a2';
            }});
            
            el.addEventListener('mouseleave', () => {{
                el.style.transform = 'perspective(1000px) rotateX(0deg) rotateY(0deg) scale3d(1, 1, 1)';
                el.style.boxShadow = '';
                el.style.borderColor = '';
            }});
        }});

        // Enter key to log in
        document.getElementById('nickname-input').addEventListener('keydown', e => {{
            if (e.key === 'Enter') {{
                grantAccess();
            }}
        }});

        // Fill nickname from local storage if available, but do not auto-login
        const storedUser = localStorage.getItem('youpass_student_id');
        if (storedUser) {{
            document.getElementById('nickname-input').value = storedUser;
        }}
        initLogin();

        // HTML5 Confetti Sparkles, Stars, and Hearts Engine
        const canvas = document.getElementById('fireworks-canvas');
        const ctx = canvas.getContext('2d');
        let animationFrameId = null;
        let particles = [];

        function resizeCanvas() {{
            canvas.width = window.innerWidth;
            canvas.height = window.innerHeight;
        }}
        window.addEventListener('resize', resizeCanvas);
        resizeCanvas();

        class Particle {{
            constructor(x, y, color) {{
                this.x = x;
                this.y = y;
                this.color = color;
                this.type = Math.random() < 0.5 ? 'star' : 'heart';
                this.angle = Math.random() * Math.PI * 2;
                this.speed = Math.random() * 5 + 2;
                this.vx = Math.cos(this.angle) * this.speed;
                this.vy = Math.sin(this.angle) * this.speed;
                this.alpha = 1;
                this.decay = Math.random() * 0.015 + 0.01;
                this.gravity = 0.06;
                this.size = Math.random() * 8 + 6;
                this.rotation = Math.random() * Math.PI;
                this.rotSpeed = (Math.random() - 0.5) * 0.1;
            }}
            update() {{
                this.x += this.vx;
                this.y += this.vy;
                this.vy += this.gravity;
                this.alpha -= this.decay;
                this.rotation += this.rotSpeed;
            }}
            draw() {{
                ctx.save();
                ctx.globalAlpha = this.alpha;
                ctx.fillStyle = this.color;
                ctx.translate(this.x, this.y);
                ctx.rotate(this.rotation);
                ctx.beginPath();
                if (this.type === 'star') {{
                    const spikes = 5;
                    let rot = Math.PI / 2 * 3;
                    let step = Math.PI / spikes;
                    ctx.moveTo(0, -this.size);
                    for (let i = 0; i < spikes; i++) {{
                        ctx.lineTo(Math.cos(rot) * this.size, Math.sin(rot) * this.size);
                        rot += step;
                        ctx.lineTo(Math.cos(rot) * (this.size/2), Math.sin(rot) * (this.size/2));
                        rot += step;
                    }}
                }} else {{
                    ctx.moveTo(0, -this.size/4);
                    ctx.bezierCurveTo(-this.size/2, -this.size, -this.size, -this.size/3, 0, this.size);
                    ctx.bezierCurveTo(this.size, -this.size/3, this.size/2, -this.size, 0, -this.size/4);
                }}
                ctx.closePath();
                ctx.fill();
                ctx.restore();
            }}
        }}

        function spawnFirework(x, y) {{
            const colors = ['#ff85a2', '#ffccd5', '#b19ffb', '#e5e0ff', '#eab308', '#ec4899', '#76e2b2'];
            const color = colors[Math.floor(Math.random() * colors.length)];
            for (let i = 0; i < 40; i++) {{
                particles.push(new Particle(x, y, color));
            }}
        }}

        function animateFireworks() {{
            ctx.clearRect(0, 0, canvas.width, canvas.height);
            
            if (Math.random() < 0.08) {{
                spawnFirework(Math.random() * canvas.width, Math.random() * canvas.height * 0.6);
            }}

            for (let i = particles.length - 1; i >= 0; i--) {{
                const p = particles[i];
                p.update();
                p.draw();
                if (p.alpha <= 0) {{
                    particles.splice(i, 1);
                }}
            }}
            animationFrameId = requestAnimationFrame(animateFireworks);
        }}

        function startFireworks() {{
            particles = [];
            animateFireworks();
        }}

        function stopFireworks() {{
            if (animationFrameId) {{
                cancelAnimationFrame(animationFrameId);
            }}
            ctx.clearRect(0, 0, canvas.width, canvas.height);
        }}

        // 9.0 Strategy Helper Script
        const strategies = {json.dumps(strategies_data, ensure_ascii=False)};
        const testType = "{test_type}";
        let activeSandboxQ = null;
        let activeTab = 'target';

        function showStrategy(qNum, event) {{
            event.stopPropagation();
            activeSandboxQ = qNum;
            const strategy = strategies[qNum - 1] || {{
                type_tip: "Focus on qualifying words and synonym changes.",
                scan_target: "Scan the passage for key terms.",
                analysis_logic: "Match the meaning to determine the correct response."
            }};
            
            document.getElementById('strategy-qnum').innerText = qNum;
            document.getElementById('strategy-tab-content-target').innerHTML = `<p style="padding: 5px 0;">${{strategy.type_tip}}</p>`;
            
            if (testType === 'listening') {{
                document.getElementById('tab-btn-synonyms').innerText = '🔊 Tape Script';
                document.getElementById('strategy-tab-content-synonyms').innerHTML = `
                    <div style="padding: 5px 0;">
                        <strong style="color: var(--accent-hover); display: block; margin-bottom: 5px;">🔊 Exact Tape Script Snippet:</strong>
                        <blockquote style="background: rgba(255, 204, 213, 0.15); border-left: 3.5px solid var(--accent); padding: 8px 12px; border-radius: 8px; font-style: italic; color: var(--text-primary); font-size: 0.85rem; line-height: 1.4;">
                            "${{strategy.scan_target}}"
                        </blockquote>
                    </div>
                `;
            }} else {{
                document.getElementById('tab-btn-synonyms').innerText = '🔄 Synonyms';
                document.getElementById('strategy-tab-content-synonyms').innerHTML = `
                    <div class="synonym-map-table">
                        <div class="synonym-row header-row">
                            <div>Keyword in Question</div>
                            <div>Expected Synonym in Text</div>
                        </div>
                        <div class="synonym-row">
                            <div style="font-weight: 700; color: var(--accent-hover);">${{strategy.scan_target}}</div>
                            <div style="color: var(--success); font-style: italic;">Check the highlighted text</div>
                        </div>
                    </div>
                `;
            }}
            
            document.getElementById('strategy-tab-content-trap').innerHTML = `<p style="padding: 5px 0;">${{strategy.analysis_logic}}</p>`;
            
            // Switch to default active tab
            switchStrategyTab(activeTab);
            
            const popover = document.getElementById('strategy-popover');
            const rect = event.target.getBoundingClientRect();
            const container = document.querySelector('.answer-pane');
            const containerRect = container.getBoundingClientRect();
            
            popover.style.display = 'block';
            
            // Calculate relative positioning inside the scrolling container
            let topPos = rect.top - containerRect.top + container.scrollTop - 90;
            let leftPos = rect.left - containerRect.left - 335;
            
            if (leftPos < 0) {{
                leftPos = rect.left - containerRect.left + 35;
            }}
            
            popover.style.top = `${{topPos}}px`;
            popover.style.left = `${{leftPos}}px`;
            
            // Auto-highlight target in passage if clue is set
            if (strategy.scan_target) {{
                highlightClueInPassage(strategy.scan_target);
            }}
        }}

        function closeStrategy() {{
            document.getElementById('strategy-popover').style.display = 'none';
            removeTempHighlights();
        }}

        function switchStrategyTab(tabName) {{
            activeTab = tabName;
            document.querySelectorAll('.strategy-tab-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelectorAll('.strategy-tab-content').forEach(content => content.classList.remove('active-content'));
            
            if (tabName === 'target') {{
                document.getElementById('tab-btn-target').classList.add('active');
                document.getElementById('strategy-tab-content-target').classList.add('active-content');
            }} else if (tabName === 'synonyms') {{
                document.getElementById('tab-btn-synonyms').classList.add('active');
                document.getElementById('strategy-tab-content-synonyms').classList.add('active-content');
            }} else if (tabName === 'trap') {{
                document.getElementById('tab-btn-trap').classList.add('active');
                document.getElementById('strategy-tab-content-trap').classList.add('active-content');
            }}
        }}

        let currentTempHighlight = null;
        function highlightClueInPassage(keywords) {{
            removeTempHighlights();
            const passagePane = document.querySelector('.content-html');
            if (!passagePane || !keywords) return;
            
            // Strip out common leading instruction phrases to get clean search keys
            let cleanKey = keywords.replace(/Scan for|Locate the phrase|Locate the word|Locate|in the/gi, '').trim();
            // Remove quotation marks if present
            cleanKey = cleanKey.replace(/^['"]|['"]$/g, '').trim();
            if (cleanKey.length < 3) return;
            
            const walker = document.createTreeWalker(passagePane, NodeFilter.SHOW_TEXT, null, false);
            let node;
            const nodesToReplace = [];
            while (node = walker.nextNode()) {{
                if (node.nodeValue.toLowerCase().includes(cleanKey.toLowerCase())) {{
                    nodesToReplace.push(node);
                }}
            }}
            
            const regex = new RegExp(`(${{escapeRegExp(cleanKey)}})`, 'gi');
            nodesToReplace.forEach(textNode => {{
                const parent = textNode.parentNode;
                if (parent && parent.className !== 'clue-highlight-temp') {{
                    const span = document.createElement('span');
                    span.className = 'clue-highlight-temp';
                    span.style.backgroundColor = 'rgba(255, 243, 128, 0.65)';
                    span.style.borderBottom = '2.5px dashed var(--accent)';
                    span.style.borderRadius = '6px';
                    span.style.padding = '1px 3px';
                    span.style.boxShadow = '0 2px 8px rgba(255, 243, 128, 0.3)';
                    
                    const origText = textNode.nodeValue;
                    const parts = origText.split(regex);
                    parts.forEach(part => {{
                        if (part.toLowerCase() === cleanKey.toLowerCase()) {{
                            const subSpan = document.createElement('span');
                            subSpan.className = 'clue-match';
                            subSpan.style.fontWeight = '700';
                            subSpan.style.color = 'var(--accent-hover)';
                            subSpan.innerText = part;
                            span.appendChild(subSpan);
                        }} else {{
                            span.appendChild(document.createTextNode(part));
                        }}
                    }});
                    parent.replaceChild(span, textNode);
                }}
            }});
        }}

        function removeTempHighlights() {{
            const temps = document.querySelectorAll('.clue-highlight-temp');
            temps.forEach(temp => {{
                const parent = temp.parentNode;
                if (parent) {{
                    parent.replaceChild(document.createTextNode(temp.textContent), temp);
                }}
            }});
        }}

        function escapeRegExp(string) {{
            return string.replace(/[.*+?^${{}}()|[\]\\#]/g, '\\$&');
        }}

        document.getElementById('strategy-sandbox-btn').onclick = function() {{
            if (activeSandboxQ) {{
                const inputEl = document.getElementById(`q-${{activeSandboxQ}}`);
                const feedbackEl = document.getElementById(`feedback-${{activeSandboxQ}}`);
                if (inputEl) {{
                    inputEl.value = '';
                    inputEl.className = 'q-input';
                    inputEl.focus();
                    playSfx('sfx-click');
                }}
                if (feedbackEl) {{
                    feedbackEl.style.display = 'none';
                }}
                closeStrategy();
            }}
        }};

        // Chống copy, chặn chuột phải, F12 và lưu trang
        document.addEventListener('contextmenu', function(e) {{
            e.preventDefault();
        }});
        document.addEventListener('keydown', function(e) {{
            if (e.key === 'F12') {{
                e.preventDefault();
                return false;
            }}
            if ((e.ctrlKey || e.metaKey) && e.shiftKey && (e.key === 'I' || e.key === 'i' || e.key === 'J' || e.key === 'j')) {{
                e.preventDefault();
                return false;
            }}
            if ((e.ctrlKey || e.metaKey) && (e.key === 'S' || e.key === 's')) {{
                e.preventDefault();
                return false;
            }}
            if ((e.ctrlKey || e.metaKey) && (e.key === 'U' || e.key === 'u')) {{
                e.preventDefault();
                return false;
            }}
        }});
        document.addEventListener('copy', function(e) {{
            e.preventDefault();
        }});
        document.addEventListener('cut', function(e) {{
            e.preventDefault();
        }});
    </script>
</body>
</html>'''

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_template)
    
    doc = Document()
    
    # Set page margins: Top=1cm, Bottom=1cm, Left=2cm, Right=1cm
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"IELTS {test_type.capitalize()} Practice - Test {idx_num}")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    parse_element_to_docx(entry_content, doc, test_folder)
                
    doc.add_page_break()
    ans_title_p = doc.add_paragraph()
    ans_title_run = ans_title_p.add_run("ANSWER KEY")
    ans_title_run.font.bold = True
    ans_title_run.font.size = Pt(14)
    ans_title_run.font.color.rgb = RGBColor(16, 185, 129)
    ans_title_p.paragraph_format.space_after = Pt(12)
    
    for a_idx, ans in enumerate(answers):
        doc.add_paragraph(f"{a_idx + 1}. {ans}")
        
    doc.save(docx_path)
    print(f"Successfully upgraded DOCX: {docx_path}")

def main():
    for test_type in ["listening", "reading"]:
        type_dir = os.path.join(DATA_DIR, test_type)
        if not os.path.exists(type_dir):
            continue
        
        folders = [f for f in os.listdir(type_dir) if os.path.isdir(os.path.join(type_dir, f)) and f.startswith("Test_")]
        
        def get_num(name):
            try:
                return int(name.split('_')[1])
            except:
                return 9999
        folders.sort(key=get_num)
        
        for folder in folders:
            idx_num = get_num(folder)
            test_folder = os.path.join(type_dir, folder)
            try:
                rebuild_files_for_test(test_folder, test_type, idx_num)
            except Exception as e:
                print(f"Error rebuilding {test_folder}: {e}")

if __name__ == '__main__':
    main()
